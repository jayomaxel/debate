"""
鐭ヨ瘑搴撶鐞嗘湇鍔?
澶勭悊鏂囨。涓婁紶銆佹枃鏈彁鍙栥€佸悜閲忕敓鎴愬拰妫€绱?
"""
import os
import json
import math
from logging_config import get_logger
from typing import List, Optional, Dict, Any
from datetime import datetime
import PyPDF2
import docx
from sqlalchemy.orm import Session
from sqlalchemy import text
import httpx

from models.document import Document
from services.config_service import ConfigService

logger = get_logger(__name__)


class KnowledgeBase:
    """鐭ヨ瘑搴撶鐞嗙被"""
    
    def __init__(self, db: Session):
        """
        鍒濆鍖栫煡璇嗗簱
        
        Args:
            db: 鏁版嵁搴撲細璇?
        """
        self.db = db
        self.upload_dir = os.getenv("UPLOAD_DIR", "uploads/documents")
        os.makedirs(self.upload_dir, exist_ok=True)

    def _is_postgresql(self) -> bool:
        return self.db.bind is not None and self.db.bind.dialect.name == "postgresql"

    def _ensure_embedding_store(self, dimension: int) -> None:
        if self._is_postgresql():
            self.db.execute(text("CREATE EXTENSION IF NOT EXISTS vector"))
            self.db.execute(text(f"""
                CREATE TABLE IF NOT EXISTS document_embeddings (
                    document_id UUID PRIMARY KEY REFERENCES documents(id) ON DELETE CASCADE,
                    embedding vector({int(dimension)}) NOT NULL,
                    embedding_json TEXT NOT NULL,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """))
            self.db.execute(text("""
                CREATE INDEX IF NOT EXISTS idx_document_embeddings_vector
                ON document_embeddings USING ivfflat (embedding vector_cosine_ops)
            """))
        else:
            self.db.execute(text("""
                CREATE TABLE IF NOT EXISTS document_embeddings (
                    document_id VARCHAR(64) PRIMARY KEY,
                    embedding_json TEXT NOT NULL,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """))
        self.db.commit()

    @staticmethod
    def _cosine_similarity(left: List[float], right: List[float]) -> float:
        if not left or not right or len(left) != len(right):
            return 0.0
        dot = sum(a * b for a, b in zip(left, right))
        left_norm = math.sqrt(sum(a * a for a in left))
        right_norm = math.sqrt(sum(b * b for b in right))
        if left_norm == 0 or right_norm == 0:
            return 0.0
        return dot / (left_norm * right_norm)
    
    async def upload_document(
        self,
        file_data: bytes,
        filename: str,
        file_type: str,
        debate_id: str
    ) -> Document:
        """
        涓婁紶鏂囨。
        
        Args:
            file_data: 鏂囦欢鏁版嵁
            filename: 鏂囦欢鍚?
            file_type: 鏂囦欢绫诲瀷锛坅pplication/pdf 鎴?application/vnd.openxmlformats-officedocument.wordprocessingml.document锛?
            debate_id: 杈╄ID
            
        Returns:
            Document瀵硅薄
        """
        try:
            # 楠岃瘉鏂囦欢绫诲瀷
            allowed_types = [
                "application/pdf",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ]
            if file_type not in allowed_types:
                raise ValueError(f"涓嶆敮鎸佺殑鏂囦欢绫诲瀷: {file_type}")
            
            # 楠岃瘉鏂囦欢澶у皬锛堟渶澶?0MB锛?
            max_size = 10 * 1024 * 1024
            if len(file_data) > max_size:
                raise ValueError("文件大小超过限制（最大10MB）")
            
            # 鐢熸垚鏂囦欢璺緞
            timestamp = datetime.utcnow().timestamp()
            safe_filename = f"{debate_id}_{timestamp}_{filename}"
            file_path = os.path.join(self.upload_dir, safe_filename)
            
            # 淇濆瓨鏂囦欢
            with open(file_path, 'wb') as f:
                f.write(file_data)
            
            logger.info(f"Document saved: {file_path}")
            
            # 鎻愬彇鏂囨湰鍐呭
            content = await self.extract_text(file_path, file_type)
            
            # 鍒涘缓鏂囨。璁板綍
            document = Document(
                debate_id=debate_id,
                filename=filename,
                file_path=file_path,
                file_type=file_type,
                content=content,
                embedding_status="pending"
            )
            
            self.db.add(document)
            self.db.commit()
            self.db.refresh(document)
            
            # 寮傛鐢熸垚鍚戦噺锛堟爣璁颁负processing锛?
            document.embedding_status = "processing"
            self.db.commit()
            
            # 鐢熸垚骞跺瓨鍌ㄥ悜閲?
            try:
                await self.generate_and_store_embeddings(document)
                document.embedding_status = "completed"
            except Exception as e:
                logger.error(f"Failed to generate embeddings: {e}", exc_info=True)
                document.embedding_status = "failed"
            
            self.db.commit()
            
            return document
            
        except Exception as e:
            logger.error(f"Failed to upload document: {e}", exc_info=True)
            raise
    
    async def extract_text(self, file_path: str, file_type: str) -> str:
        """
        浠庢枃浠朵腑鎻愬彇鏂囨湰鍐呭
        
        Args:
            file_path: 鏂囦欢璺緞
            file_type: 鏂囦欢绫诲瀷
            
        Returns:
            鎻愬彇鐨勬枃鏈唴瀹?
        """
        try:
            if file_type == "application/pdf":
                return await self._extract_text_from_pdf(file_path)
            elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                return await self._extract_text_from_docx(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        except Exception as e:
            logger.error(f"Failed to extract text: {e}", exc_info=True)
            raise
    
    async def _extract_text_from_pdf(self, file_path: str) -> str:
        """浠嶱DF鏂囦欢鎻愬彇鏂囨湰"""
        try:
            text_content = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text:
                        text_content.append(text)
            
            full_text = "\n\n".join(text_content)
            logger.info(f"Extracted {len(full_text)} characters from PDF")
            
            return full_text
            
        except Exception as e:
            logger.error(f"Failed to extract text from PDF: {e}", exc_info=True)
            raise
    
    async def _extract_text_from_docx(self, file_path: str) -> str:
        """浠嶹ord鏂囨。鎻愬彇鏂囨湰"""
        try:
            doc = docx.Document(file_path)
            
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            full_text = "\n\n".join(text_content)
            logger.info(f"Extracted {len(full_text)} characters from DOCX")
            
            return full_text
            
        except Exception as e:
            logger.error(f"Failed to extract text from DOCX: {e}", exc_info=True)
            raise
    
    async def generate_embeddings(self, text: str) -> List[float]:
        """
        鐢熸垚鏂囨湰鐨勮涔夊悜閲?
        
        Args:
            text: 鏂囨湰鍐呭
            
        Returns:
            鍚戦噺鍒楄〃
        """
        try:
            # 鑾峰彇OpenAI閰嶇疆
            config_service = ConfigService(self.db)
            model_config = await config_service.get_model_config()
            
            if not model_config or not model_config.api_key:
                raise ValueError("OpenAI API key not configured")
            
            # 鍒嗗潡澶勭悊闀挎枃鏈紙OpenAI embeddings API闄愬埗8191 tokens锛?
            chunks = self._split_text_into_chunks(text, max_length=8000)
            
            all_embeddings = []
            
            async with httpx.AsyncClient(timeout=30.0) as client:
                for chunk in chunks:
                    response = await client.post(
                        f"{model_config.api_endpoint.rsplit('/', 1)[0]}/embeddings",
                        headers={
                            "Authorization": f"Bearer {model_config.api_key}",
                            "Content-Type": "application/json"
                        },
                        json={
                            "model": "text-embedding-ada-002",
                            "input": chunk
                        }
                    )
                    
                    if response.status_code == 200:
                        data = response.json()
                        embedding = data["data"][0]["embedding"]
                        all_embeddings.append(embedding)
                    else:
                        error_msg = f"Embeddings API error: {response.status_code} - {response.text}"
                        logger.error(error_msg)
                        raise Exception(error_msg)
            
            # 濡傛灉鏈夊涓猚hunk锛岃繑鍥炵涓€涓紙鎴栧彲浠ヨ绠楀钩鍧囧€硷級
            return all_embeddings[0] if all_embeddings else []
            
        except Exception as e:
            logger.error(f"Failed to generate embeddings: {e}", exc_info=True)
            raise
    
    def _split_text_into_chunks(self, text: str, max_length: int = 8000) -> List[str]:
        """
        灏嗛暱鏂囨湰鍒嗗壊鎴愬涓潡
        
        Args:
            text: 鏂囨湰鍐呭
            max_length: 姣忓潡鐨勬渶澶ч暱搴?
            
        Returns:
            鏂囨湰鍧楀垪琛?
        """
        if len(text) <= max_length:
            return [text]
        
        chunks = []
        current_chunk = ""
        
        # 鎸夋钀藉垎鍓?
        paragraphs = text.split("\n\n")
        
        for paragraph in paragraphs:
            if len(current_chunk) + len(paragraph) + 2 <= max_length:
                current_chunk += paragraph + "\n\n"
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = paragraph + "\n\n"
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return chunks
    
    async def generate_and_store_embeddings(self, document: Document) -> None:
        """
        鐢熸垚骞跺瓨鍌ㄦ枃妗ｇ殑鍚戦噺鍒版暟鎹簱
        
        Args:
            document: 鏂囨。瀵硅薄
        """
        try:
            if not document.content:
                logger.warning(f"Document {document.id} has no content")
                return
            
            # 鐢熸垚鍚戦噺
            embeddings = await self.generate_embeddings(document.content)
            
            if not embeddings:
                logger.warning(f"No embeddings generated for document {document.id}")
                return

            self._ensure_embedding_store(len(embeddings))
            embedding_json = json.dumps(embeddings)

            if self._is_postgresql():
                vector_str = "[" + ",".join(str(x) for x in embeddings) + "]"
                self.db.execute(
                    text("""
                        INSERT INTO document_embeddings (document_id, embedding, embedding_json, updated_at)
                        VALUES (:document_id, CAST(:embedding AS vector), :embedding_json, CURRENT_TIMESTAMP)
                        ON CONFLICT (document_id) DO UPDATE SET
                            embedding = EXCLUDED.embedding,
                            embedding_json = EXCLUDED.embedding_json,
                            updated_at = CURRENT_TIMESTAMP
                    """),
                    {
                        "document_id": str(document.id),
                        "embedding": vector_str,
                        "embedding_json": embedding_json,
                    },
                )
            else:
                self.db.execute(
                    text("""
                        INSERT OR REPLACE INTO document_embeddings
                            (document_id, embedding_json, updated_at)
                        VALUES (:document_id, :embedding_json, CURRENT_TIMESTAMP)
                    """),
                    {
                        "document_id": str(document.id),
                        "embedding_json": embedding_json,
                    },
                )

            self.db.commit()
            
            
        except Exception as e:
            logger.error(f"Failed to generate and store embeddings: {e}", exc_info=True)
            raise
    
    async def search_relevant_content(
        self,
        query: str,
        debate_id: Optional[str] = None,
        top_k: int = 5
    ) -> List[Dict[str, Any]]:
        """
        鎼滅储鐩稿叧鍐呭
        
        Args:
            query: 鏌ヨ鏂囨湰
            debate_id: 杈╄ID锛堝彲閫夛紝鐢ㄤ簬闄愬埗鎼滅储鑼冨洿锛?
            top_k: 杩斿洖缁撴灉鏁伴噺
            
        Returns:
            鐩稿叧鏂囨。鍒楄〃
        """
        try:
            # 鐢熸垚鏌ヨ鍚戦噺
            query_embedding = await self.generate_embeddings(query)
            
            if not query_embedding:
                logger.warning("Failed to generate query embedding")
                return []

            self._ensure_embedding_store(len(query_embedding))

            if self._is_postgresql():
                vector_str = "[" + ",".join(str(x) for x in query_embedding) + "]"
                sql = """
                    SELECT
                        d.id AS document_id,
                        d.filename,
                        d.content,
                        1 - (e.embedding <=> CAST(:query_vector AS vector)) AS relevance_score
                    FROM document_embeddings e
                    JOIN documents d ON d.id = e.document_id
                    WHERE d.embedding_status = 'completed'
                """
                params: Dict[str, Any] = {
                    "query_vector": vector_str,
                    "limit": top_k,
                }
                if debate_id:
                    sql += " AND d.debate_id = CAST(:debate_id AS UUID)"
                    params["debate_id"] = debate_id
                sql += " ORDER BY e.embedding <=> CAST(:query_vector AS vector) LIMIT :limit"

                rows = self.db.execute(text(sql), params).fetchall()
                return [
                    {
                        "document_id": str(row.document_id),
                        "filename": row.filename,
                        "content": row.content,
                        "relevance_score": float(row.relevance_score),
                    }
                    for row in rows
                ]

            sql = """
                SELECT d.id AS document_id, d.filename, d.content, e.embedding_json
                FROM document_embeddings e
                JOIN documents d ON CAST(d.id AS TEXT) = e.document_id
                WHERE d.embedding_status = 'completed'
            """
            params = {}
            if debate_id:
                sql += " AND CAST(d.debate_id AS TEXT) = :debate_id"
                params["debate_id"] = str(debate_id)

            rows = self.db.execute(text(sql), params).fetchall()
            results = []
            for row in rows:
                try:
                    document_embedding = json.loads(row.embedding_json)
                except (TypeError, json.JSONDecodeError):
                    continue
                score = self._cosine_similarity(query_embedding, document_embedding)
                results.append({
                    "document_id": str(row.document_id),
                    "filename": row.filename,
                    "content": row.content,
                    "relevance_score": score,
                })

            results.sort(key=lambda item: item["relevance_score"], reverse=True)
            return results[:top_k]
            
            
        except Exception as e:
            logger.error(f"Failed to search relevant content: {e}", exc_info=True)
            return []
    
    async def delete_document(self, document_id: str) -> bool:
        """
        鍒犻櫎鏂囨。
        
        Args:
            document_id: 鏂囨。ID
            
        Returns:
            鏄惁鍒犻櫎鎴愬姛
        """
        try:
            document = self.db.query(Document).filter(Document.id == document_id).first()
            
            if not document:
                logger.warning(f"Document {document_id} not found")
                return False
            
            # 鍒犻櫎鏂囦欢
            if os.path.exists(document.file_path):
                os.remove(document.file_path)
                logger.info(f"Deleted file: {document.file_path}")

            self.db.execute(
                text("DELETE FROM document_embeddings WHERE document_id = :document_id"),
                {"document_id": str(document_id)}
            )
            
            # 鍒犻櫎鏁版嵁搴撹褰?
            self.db.delete(document)
            self.db.commit()
            
            logger.info(f"Deleted document: {document_id}")
            return True
            
        except Exception as e:
            logger.error(f"Failed to delete document: {e}", exc_info=True)
            self.db.rollback()
            return False
    
    def get_documents(self, debate_id: str) -> List[Document]:
        """
        鑾峰彇杈╄鐨勬墍鏈夋枃妗?
        
        Args:
            debate_id: 杈╄ID
            
        Returns:
            鏂囨。鍒楄〃
        """
        try:
            documents = self.db.query(Document).filter(
                Document.debate_id == debate_id
            ).order_by(Document.uploaded_at.desc()).all()
            
            return documents
            
        except Exception as e:
            logger.error(f"Failed to get documents: {e}", exc_info=True)
            return []
    
    def get_document_by_id(self, document_id: str) -> Optional[Document]:
        """
        鏍规嵁ID鑾峰彇鏂囨。
        
        Args:
            document_id: 鏂囨。ID
            
        Returns:
            鏂囨。瀵硅薄
        """
        try:
            document = self.db.query(Document).filter(Document.id == document_id).first()
            return document
        except Exception as e:
            logger.error(f"Failed to get document: {e}", exc_info=True)
            return None

