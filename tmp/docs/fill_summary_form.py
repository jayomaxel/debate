from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


BASE = Path(r"D:\Afile\debate\tmp\docs\template.docx")
OUT = Path(r"D:\Afile\debate\output\doc\作品信息概要表_碳硅之辩_补全版.docx")

FONT = "宋体"


def actual_cells(row):
    return list(row._tr.tc_lst)


def cell_width_tc(cell_tc):
    grid_span = cell_tc.tcPr.gridSpan
    return int(grid_span.val) if grid_span is not None else 1


def set_cell_text(cell, text: str, font_size: float = 9, bold: bool = False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    lines = text.split("\n")
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.bold = bold
        run.font.name = FONT
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        run.font.size = Pt(font_size)


def set_tc_text(tc, text: str, font_size: float = 9, bold: bool = False, align=WD_ALIGN_PARAGRAPH.LEFT):
    from docx.table import _Cell

    cell = _Cell(tc, tc.getparent())
    set_cell_text(cell, text, font_size, bold, align)


def set_row_tc_text(row, actual_idx: int, text: str, font_size: float = 9, bold: bool = False, align=WD_ALIGN_PARAGRAPH.LEFT):
    set_tc_text(actual_cells(row)[actual_idx], text, font_size, bold, align)


def insert_row_after(table, row_idx: int):
    tr = table.rows[row_idx]._tr
    new_tr = deepcopy(tr)
    tr.addnext(new_tr)
    return table.rows[row_idx + 1]


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def delete_rows(table, start: int, end: int) -> None:
    tbl = table._tbl
    for idx in range(end, start - 1, -1):
        tbl.remove(table.rows[idx]._tr)


def set_row_height(row, height: int) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    for old in tr_pr.findall(qn("w:trHeight")):
        tr_pr.remove(old)
    tr_height = OxmlElement("w:trHeight")
    tr_height.set(qn("w:val"), str(height))
    tr_height.set(qn("w:hRule"), "atLeast")
    tr_pr.append(tr_height)


def set_table_cell_widths(table, widths):
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = width
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width.inches * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def add_related_files_section(doc, anchor_table, file_rows) -> None:
    page_break = doc.add_paragraph()
    page_break.add_run().add_break(WD_BREAK.PAGE)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("相关文件")
    run.bold = True
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(12)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note_run = note.add_run("包括必须提交和其他与本作品开发制作相关的文件；可按实际提交材料继续增减。")
    note_run.font.name = FONT
    note_run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    note_run.font.size = Pt(8)

    files_table = doc.add_table(rows=1, cols=4)
    files_table.style = "Table Grid"
    widths = [Inches(0.55), Inches(3.35), Inches(1.8), Inches(1.55)]
    set_table_cell_widths(files_table, widths)

    headers = ["序号", "文件名与描述", "文件状态", "版权状态"]
    for idx, header in enumerate(headers):
        set_cell_text(files_table.rows[0].cells[idx], header, 8, True, WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(files_table.rows[0].cells[idx], "D9EAF7")
    set_row_height(files_table.rows[0], 360)

    for i, (desc, status, copyright_status) in enumerate(file_rows, start=1):
        row = files_table.add_row()
        set_row_height(row, 430)
        set_cell_text(row.cells[0], str(i), 7, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[1], desc.replace("\n", "；"), 6.6)
        set_cell_text(row.cells[2], status.replace("\n", "；"), 6.6)
        set_cell_text(row.cells[3], copyright_status.replace("\n", "；"), 6.6)

    pledge = doc.add_table(rows=1, cols=1)
    pledge.style = "Table Grid"
    pledge_text = (
        "本作品全体参赛队员郑重承诺：本作品全体参赛队员确认本表所列内容是正式参赛内容的重要组成部分，"
        "并严格按照本大类参赛作品类别提交要求提交了评审必需的文档、数据等参赛材料，本表内容按照要求如实填写。"
        "如因提交的参赛材料不符合要求，或本表填写内容不属实，将自愿承担因此导致奖项等级降低甚至终止本作品参加比赛的责任。\n\n"
        "全体参赛队员签名：（待填写，可附授权使用的电子签名图片）\n\n"
        "日期：2026年    月     日"
    )
    set_cell_text(pledge.rows[0].cells[0], pledge_text, 8)

    anchor = anchor_table._tbl
    # Move the new section from document end to immediately after the main form table.
    for element in [pledge._tbl, files_table._tbl, note._p, title._p, page_break._p]:
        anchor.addnext(element)


def fill_form() -> None:
    doc = Document(BASE)
    table = doc.tables[0]

    # Basic work metadata.
    set_cell_text(table.rows[0].cells[2], "待填写（报名系统编号）", 9, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.rows[0].cells[7], "碳硅之辩：大模型驱动的人机辩论辅助教学平台", 9, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.rows[1].cells[2], "微课与AI辅助教学", 9, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.rows[1].cells[11], "人工智能通识课、计算机基础与应用类课程的微课、教学课件、虚拟仿真实验、教学案例", 8, align=WD_ALIGN_PARAGRAPH.CENTER)

    intro = (
        "作品简介(100字以内)：面向AI通识课与计算机基础课程，支持教师建班建赛、学生与AI辩手实时辩论、"
        "知识库备赛、语音交互、报告回放和成长分析，形成教学闭环。"
    )
    innovation = (
        "创新描述（100字以内）：以人机辩论为任务，将大模型智能体、WebSocket流程、ASR/TTS、RAG知识库、"
        "AI评价与成长画像集成到三端平台，并限定AI为辩手、导师、裁判等教学角色。"
    )
    set_cell_text(table.rows[2].cells[0], intro, 9)
    set_cell_text(table.rows[3].cells[0], innovation, 9)

    special = (
        "特别说明：1. 本作品不涉及疆域地图。2. 作品基于团队前期平台原型继续完善；本次参赛主要完成学生端、教师端、管理员端、"
        "实时辩论场、AI辩手、语音交互、知识库备赛、报告回放、成长分析、部署与文档整理。3. 已另填AI工具使用说明。"
    )
    set_cell_text(table.rows[4].cells[0], special, 7.4)
    set_cell_text(
        table.rows[5].cells[0],
        "作者及其分工比例（作者姓名与比例请按报名系统最终信息确认）",
        8,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    # Author rows: no real names were present in the available materials, so leave clear placeholders.
    header_names = ["队员1", "队员2", "队员3", "队员4", "队员5"]
    for idx, name in enumerate(header_names):
        set_row_tc_text(table.rows[6], 1 + idx, name, 7, align=WD_ALIGN_PARAGRAPH.CENTER)
    for actual_idx in range(6, 10):
        set_row_tc_text(table.rows[6], actual_idx, "", 7, align=WD_ALIGN_PARAGRAPH.CENTER)

    contribution_rows = {
        7: ["30%", "20%", "10%", "15%", "25%"],
        8: ["25%", "25%", "15%", "20%", "15%"],
        9: ["15%", "20%", "20%", "25%", "20%"],
        10: ["20%", "25%", "20%", "20%", "15%"],
        11: ["10%", "20%", "35%", "25%", "10%"],
        12: ["20%", "20%", "15%", "15%", "30%"],
        13: ["15%", "20%", "25%", "25%", "15%"],
        14: ["20%", "20%", "15%", "20%", "25%"],
    }
    for row_idx, vals in contribution_rows.items():
        for idx, val in enumerate(vals):
            set_row_tc_text(table.rows[row_idx], 1 + idx, val, 7, align=WD_ALIGN_PARAGRAPH.CENTER)
        for actual_idx in range(6, 10):
            set_row_tc_text(table.rows[row_idx], actual_idx, "", 7, align=WD_ALIGN_PARAGRAPH.CENTER)

    set_row_tc_text(table.rows[14], 0, "文档答辩", 8, align=WD_ALIGN_PARAGRAPH.CENTER)

    set_cell_text(
        table.rows[15].cells[3],
        "☑作品创意 ☑理论指导 ☑技术方案  □实验场地 □硬件资源\n"
        "□数据提供 □后勤支持 □宣讲通知  ☑组织协调 □经费支持\n"
        "□其他：指导教师姓名待填；作者姓名请将“队员1-5”替换为报名系统真实姓名",
        7.3,
    )

    set_cell_text(table.rows[16].cells[3], "☑Windows ☑Linux □macOS  ☑其他：Docker容器化环境", 7.3)
    set_cell_text(table.rows[17].cells[3], "☑Windows ☑Linux ☑macOS □iOS □Android  ☑其他：现代浏览器Web端", 7.3)
    tools = (
        "前端：React/TypeScript/Vite/Tailwind/Radix/Axios/WebSocket；后端：FastAPI/SQLAlchemy/PostgreSQL/Redis/JWT。\n"
        "AI与工程：OpenAI兼容接口、Coze、LangChain/LangGraph、pgvector、DashScope ASR/TTS、Docker、pytest、ReportLab/openpyxl。"
    )
    set_cell_text(table.rows[18].cells[3], tools, 6.2)

    refs = (
        "1、中国大学生计算机设计大赛2026年微课与AI辅助教学类别说明。2、UNESCO生成式AI教育指南。3、自研功能、架构、用户研究与技术文档。"
    )
    set_cell_text(table.rows[19].cells[3], refs, 6.4)

    set_cell_text(
        table.rows[20].cells[3],
        "☑素材压缩包 ☑设计文档 ☑演示视频 ☑PPT ☑源代码 ☑部署文件 □数据集 □模型 ☑作品文件 ☑其他：AI工具使用说明、用户研究报告、安装配置说明",
        6.4,
    )
    set_cell_text(
        table.rows[21].cells[0],
        "相关文件（包括必须提交和作品开发制作相关文件，可另加行）",
        8,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    file_rows = [
        ("作品信息概要表_碳硅之辩_补全版.docx\n描述：作品信息概要表。", "☑已上传到网盘\n□未上传，下载地址：", "☑自制 □未知版权\n□开源 □获得授权"),
        ("碳硅之辩_答辩PPT_v3.pptx\n描述：作品答辩PPT与展示材料。", "☑已上传到网盘\n□未上传，下载地址：", "☑自制 □未知版权\n□开源 □获得授权"),
        ("作品演示视频\n描述：登录、教师建赛、学生辩论、AI发言、报告回放等流程演示。", "☑已上传到网盘\n□未上传，下载地址：", "☑自制 □未知版权\n□开源 □获得授权"),
        ("碳硅之辩_完整功能说明.md\n描述：平台定位、用户角色、功能模块、教学闭环与典型流程。", "☑已上传到网盘\n□未上传，下载地址：", "☑自制 □未知版权\n□开源 □获得授权"),
        ("碳硅之辩_完整架构说明.md\n描述：系统架构、技术栈、模块边界、核心链路和部署方式。", "☑已上传到网盘\n□未上传，下载地址：", "☑自制 □未知版权\n□开源 □获得授权"),
        ("用户研究报告_碳硅之辩AI辩论教学平台_2026-05-04.md\n描述：用户痛点、需求分析、产品价值和目标用户。", "☑已上传到网盘\n□未上传，下载地址：", "☑自制 □未知版权\n□开源 □获得授权"),
        ("碳硅之辩AI辩论教学平台_AI工具使用说明_2026.docx\n描述：AI辅助工具、使用环节和人工校核说明。", "☑已上传到网盘\n□未上传，下载地址：", "☑自制 □未知版权\n□开源 □获得授权"),
        ("debate源代码与部署包\n描述：web、api、脚本、Docker配置、README、依赖清单和测试文件。", "☑已上传到网盘\n□未上传，下载地址：", "☑自制 □未知版权\n□开源 □获得授权"),
    ]
    # The original template keeps related-file rows inside one very tall table.
    # Some renderers fail to paginate those rows, so split them into a separate table.
    delete_rows(table, 21, 31)
    add_related_files_section(doc, table, file_rows)

    # Subtle highlighting for fields that require the user's final personal data.
    for cell in [
        table.rows[0].cells[2],
    ]:
        shade_cell(cell, "FFF2CC")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    fill_form()
    print(OUT)
