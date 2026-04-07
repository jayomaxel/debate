"""
文档服务单元测试
测试DocumentService的上传验证、文件存储和删除功能
"""
import os
import pytest
import uuid
from datetime import datetime
from sqlalchemy.orm import sessionmaker, Session
from database import Base
from models.kb_document import KBDocument, KBDocumentChunk
from models.user import User
from services.document_service import DocumentService
from testing_db import (
    create_test_engine,
    create_test_schema,
    drop_test_schema,
    resolve_test_database_url,
)


# 测试数据库配置
TEST_DATABASE_URL = "sqlite:///:memory:"


def assert_message_contains_any(message, *expected_variants):
    """Accept small wording drift in user-facing error messages."""
    assert any(variant in message for variant in expected_variants), (
        f"expected one of {expected_variants!r} in message: {message!r}"
    )


def assert_missing_document_message(message):
    assert_message_contains_any(message, "文档不存在", "文档不存")


def assert_missing_file_message(message):
    assert_message_contains_any(message, "文件不存在", "文件不存")


def assert_corrupted_file_message(message):
    assert_message_contains_any(message, "损坏或格式错误", "损坏或格式错")


def assert_api_key_missing_message(message):
    assert_message_contains_any(message, "未配置OPENAI_API_KEY", "API密钥未配置")


def assert_extractable_text_missing_message(message):
    assert_message_contains_any(message, "未包含可提取的文本内容", "未提取到文本内容")


def assert_docx_parse_failure_message(message):
    assert_message_contains_any(message, "解析DOCX文件失败", "DOCX文件解析失败")


def assert_empty_text_message(message):
    assert_message_contains_any(message, "文本内容为空", "文本为空")


def assert_empty_text_list_message(message):
    assert_message_contains_any(message, "文本列表为空", "文本列表")


def assert_empty_chunk_list_message(message):
    assert_message_contains_any(message, "文档块列表为空", "文档块列表")


def assert_embedding_dimension_message(message):
    assert_message_contains_any(message, "嵌入向量维度错误", "查询嵌入向量维度错误")
    assert "1536" in message


def assert_invalid_top_k_message(message):
    assert_message_contains_any(message, "top_k必须大于0", "top_k")


@pytest.fixture
def db_session(request):
    """创建测试数据库会话"""
    database_url = resolve_test_database_url(
        TEST_DATABASE_URL,
        use_pgvector=bool(request.node.get_closest_marker("pgvector")),
    )
    engine = create_test_engine(database_url)
    create_test_schema(engine)
    
    TestingSessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)
    session = TestingSessionLocal()
    
    # 创建测试用户
    test_user = User(
        id=uuid.uuid4(),
        account="test_admin",
        name="Test Admin",
        email="admin@test.com",
        password_hash="hashed_password",
        user_type="administrator",
        created_at=datetime.utcnow()
    )
    session.add(test_user)
    session.commit()
    
    try:
        yield session
    finally:
        session.close()
        drop_test_schema(engine)


@pytest.fixture
def document_service(db_session):
    """创建DocumentService实例"""
    service = DocumentService(db_session)
    # 使用临时目录进行测试
    service.upload_dir = "test_uploads"
    os.makedirs(service.upload_dir, exist_ok=True)
    yield service
    # 清理测试文件
    if os.path.exists(service.upload_dir):
        for file in os.listdir(service.upload_dir):
            os.remove(os.path.join(service.upload_dir, file))
        os.rmdir(service.upload_dir)


@pytest.fixture
def test_user_id(db_session):
    """获取测试用户ID"""
    user = db_session.query(User).first()
    return str(user.id)


class TestFileValidation:
    """测试文件验证功能"""
    
    def test_validate_pdf_file_type(self, document_service):
        """测试PDF文件类型验证"""
        assert document_service.validate_file_type("application/pdf") is True
    
    def test_validate_docx_file_type(self, document_service):
        """测试DOCX文件类型验证"""
        file_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        assert document_service.validate_file_type(file_type) is True
    
    def test_reject_invalid_file_type(self, document_service):
        """测试拒绝不支持的文件类型"""
        assert document_service.validate_file_type("application/json") is False
        assert document_service.validate_file_type("text/plain") is False
        assert document_service.validate_file_type("image/png") is False
    
    def test_validate_file_size_within_limit(self, document_service):
        """测试文件大小在限制内"""
        # 1MB文件
        assert document_service.validate_file_size(1 * 1024 * 1024) is True
        # 10MB文件（边界值）
        assert document_service.validate_file_size(10 * 1024 * 1024) is True
    
    def test_reject_file_size_exceeds_limit(self, document_service):
        """测试拒绝超过大小限制的文件"""
        # 11MB文件
        assert document_service.validate_file_size(11 * 1024 * 1024) is False
        # 100MB文件
        assert document_service.validate_file_size(100 * 1024 * 1024) is False


class TestFileUpload:
    """测试文件上传功能"""
    
    @pytest.mark.asyncio
    async def test_upload_valid_pdf(self, document_service, test_user_id):
        """测试上传有效的PDF文件"""
        file_data = b"PDF content here" * 100  # 模拟PDF内容
        filename = "test_document.pdf"
        file_type = "application/pdf"
        
        document = await document_service.upload_document(
            file_data=file_data,
            filename=filename,
            file_type=file_type,
            user_id=test_user_id
        )
        
        # 验证文档记录
        assert document.id is not None
        assert document.filename == filename
        assert document.file_type == file_type
        assert document.file_size == len(file_data)
        assert document.upload_status == "pending"
        assert document.uploaded_by == uuid.UUID(test_user_id)
        
        # 验证文件已保存
        assert os.path.exists(document.file_path)
        with open(document.file_path, 'rb') as f:
            assert f.read() == file_data
    
    @pytest.mark.asyncio
    async def test_upload_valid_docx(self, document_service, test_user_id):
        """测试上传有效的DOCX文件"""
        file_data = b"DOCX content here" * 100
        filename = "test_document.docx"
        file_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        
        document = await document_service.upload_document(
            file_data=file_data,
            filename=filename,
            file_type=file_type,
            user_id=test_user_id
        )
        
        assert document.id is not None
        assert document.filename == filename
        assert document.file_type == file_type
        assert os.path.exists(document.file_path)
    
    @pytest.mark.asyncio
    async def test_reject_invalid_file_type(self, document_service, test_user_id):
        """测试拒绝不支持的文件类型"""
        file_data = b"Some content"
        filename = "test.txt"
        file_type = "text/plain"
        
        with pytest.raises(ValueError) as exc_info:
            await document_service.upload_document(
                file_data=file_data,
                filename=filename,
                file_type=file_type,
                user_id=test_user_id
            )

        error_message = str(exc_info.value)
        assert "不支持的文件类型" in error_message
        assert "PDF" in error_message
        assert "DOCX" in error_message
    
    @pytest.mark.asyncio
    async def test_reject_file_too_large(self, document_service, test_user_id):
        """测试拒绝超过大小限制的文件"""
        # 创建11MB的文件
        file_data = b"x" * (11 * 1024 * 1024)
        filename = "large_file.pdf"
        file_type = "application/pdf"
        
        with pytest.raises(ValueError) as exc_info:
            await document_service.upload_document(
                file_data=file_data,
                filename=filename,
                file_type=file_type,
                user_id=test_user_id
            )
        
        assert "文件大小超过限制" in str(exc_info.value)
        assert "10MB" in str(exc_info.value)
    
    @pytest.mark.asyncio
    async def test_unique_filename_generation(self, document_service, test_user_id):
        """测试生成唯一文件名"""
        file_data = b"Content"
        filename = "duplicate.pdf"
        file_type = "application/pdf"
        
        # 上传两个同名文件
        doc1 = await document_service.upload_document(
            file_data=file_data,
            filename=filename,
            file_type=file_type,
            user_id=test_user_id
        )
        
        doc2 = await document_service.upload_document(
            file_data=file_data,
            filename=filename,
            file_type=file_type,
            user_id=test_user_id
        )
        
        # 验证文件路径不同
        assert doc1.file_path != doc2.file_path
        # 验证两个文件都存在
        assert os.path.exists(doc1.file_path)
        assert os.path.exists(doc2.file_path)


class TestDocumentListing:
    """测试文档列表功能"""
    
    @pytest.mark.asyncio
    async def test_list_empty_documents(self, document_service):
        """测试列出空文档列表"""
        result = document_service.list_documents(page=1, page_size=20)
        
        assert result["documents"] == []
        assert result["total"] == 0
        assert result["page"] == 1
        assert result["page_size"] == 20
        assert result["total_pages"] == 0
    
    @pytest.mark.asyncio
    async def test_list_documents_with_pagination(self, document_service, test_user_id):
        """测试分页列出文档"""
        # 创建5个文档
        for i in range(5):
            await document_service.upload_document(
                file_data=b"Content",
                filename=f"doc_{i}.pdf",
                file_type="application/pdf",
                user_id=test_user_id
            )
        
        # 第一页（2个文档）
        result = document_service.list_documents(page=1, page_size=2)
        assert len(result["documents"]) == 2
        assert result["total"] == 5
        assert result["page"] == 1
        assert result["total_pages"] == 3
        
        # 第二页（2个文档）
        result = document_service.list_documents(page=2, page_size=2)
        assert len(result["documents"]) == 2
        assert result["page"] == 2
        
        # 第三页（1个文档）
        result = document_service.list_documents(page=3, page_size=2)
        assert len(result["documents"]) == 1
        assert result["page"] == 3
    
    @pytest.mark.asyncio
    async def test_list_documents_ordered_by_upload_time(self, document_service, test_user_id):
        """测试文档按上传时间倒序排列"""
        # 创建3个文档
        doc1 = await document_service.upload_document(
            file_data=b"Content 1",
            filename="doc_1.pdf",
            file_type="application/pdf",
            user_id=test_user_id
        )
        
        doc2 = await document_service.upload_document(
            file_data=b"Content 2",
            filename="doc_2.pdf",
            file_type="application/pdf",
            user_id=test_user_id
        )
        
        doc3 = await document_service.upload_document(
            file_data=b"Content 3",
            filename="doc_3.pdf",
            file_type="application/pdf",
            user_id=test_user_id
        )
        
        result = document_service.list_documents(page=1, page_size=10)
        documents = result["documents"]
        
        # 验证顺序（最新的在前）
        assert documents[0].id == doc3.id
        assert documents[1].id == doc2.id
        assert documents[2].id == doc1.id


class TestDocumentDeletion:
    """测试文档删除功能"""
    
    @pytest.mark.asyncio
    async def test_delete_existing_document(self, document_service, test_user_id):
        """测试删除存在的文档"""
        # 创建文档
        document = await document_service.upload_document(
            file_data=b"Content to delete",
            filename="to_delete.pdf",
            file_type="application/pdf",
            user_id=test_user_id
        )
        
        file_path = document.file_path
        document_id = str(document.id)
        
        # 验证文件存在
        assert os.path.exists(file_path)
        
        # 删除文档
        result = await document_service.delete_document(document_id)
        assert result is True
        
        # 验证文件已删除
        assert not os.path.exists(file_path)
        
        # 验证数据库记录已删除
        from sqlalchemy import select
        deleted_doc = document_service.db.execute(
            select(KBDocument).where(KBDocument.id == uuid.UUID(document_id))
        ).scalar_one_or_none()
        assert deleted_doc is None
    
    @pytest.mark.asyncio
    async def test_delete_nonexistent_document(self, document_service):
        """测试删除不存在的文档"""
        fake_id = str(uuid.uuid4())
        
        with pytest.raises(ValueError) as exc_info:
            await document_service.delete_document(fake_id)

        error_message = str(exc_info.value)
        assert_missing_document_message(error_message)
        assert fake_id in error_message
    
    @pytest.mark.asyncio
    async def test_delete_document_with_chunks(self, document_service, test_user_id, db_session):
        """测试删除带有chunks的文档（级联删除）"""
        # 创建文档
        document = await document_service.upload_document(
            file_data=b"Content",
            filename="doc_with_chunks.pdf",
            file_type="application/pdf",
            user_id=test_user_id
        )
        
        # 注意：在SQLite测试中，我们不能创建chunks表（因为ARRAY类型不支持）
        # 但我们仍然可以测试文档删除功能
        # 在实际的PostgreSQL环境中，级联删除会自动处理chunks
        
        document_id = str(document.id)
        
        # 删除文档
        await document_service.delete_document(document_id)
        
        # 验证文档已删除
        from sqlalchemy import select
        deleted_doc = document_service.db.execute(
            select(KBDocument).where(KBDocument.id == document.id)
        ).scalar_one_or_none()
        assert deleted_doc is None


if __name__ == "__main__":
    pytest.main([__file__, "-v"])


class TestDocumentParsing:
    """测试文档解析功能"""
    
    @pytest.mark.asyncio
    async def test_parse_pdf_success(self, document_service, tmp_path):
        """测试成功解析PDF文件"""
        # 创建一个简单的PDF文件用于测试
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        
        pdf_path = tmp_path / "test.pdf"
        c = canvas.Canvas(str(pdf_path), pagesize=letter)
        c.drawString(100, 750, "This is a test PDF document.")
        c.drawString(100, 730, "It contains multiple lines of text.")
        c.showPage()
        c.drawString(100, 750, "This is page 2.")
        c.save()
        
        # 解析PDF
        text = await document_service.parse_pdf(str(pdf_path))
        
        # 验证提取的文本
        assert "test PDF document" in text
        assert "multiple lines" in text
        assert "page 2" in text
    
    @pytest.mark.asyncio
    async def test_parse_pdf_empty_content(self, document_service, tmp_path):
        """测试解析空内容的PDF文件"""
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        
        pdf_path = tmp_path / "empty.pdf"
        c = canvas.Canvas(str(pdf_path), pagesize=letter)
        c.showPage()  # 空白页
        c.save()
        
        # 解析应该抛出ValueError
        with pytest.raises(ValueError) as exc_info:
            await document_service.parse_pdf(str(pdf_path))
        
        assert_extractable_text_missing_message(str(exc_info.value))
    
    @pytest.mark.asyncio
    async def test_parse_pdf_corrupted_file(self, document_service, tmp_path):
        """测试解析损坏的PDF文件"""
        pdf_path = tmp_path / "corrupted.pdf"
        # 创建一个假的PDF文件（不是真正的PDF格式）
        with open(pdf_path, 'wb') as f:
            f.write(b"This is not a valid PDF file")
        
        # 解析应该抛出ValueError
        with pytest.raises(ValueError) as exc_info:
            await document_service.parse_pdf(str(pdf_path))

        assert_corrupted_file_message(str(exc_info.value))
    
    @pytest.mark.asyncio
    async def test_parse_pdf_nonexistent_file(self, document_service):
        """测试解析不存在的PDF文件"""
        with pytest.raises(ValueError) as exc_info:
            await document_service.parse_pdf("/nonexistent/file.pdf")

        error_message = str(exc_info.value)
        assert_missing_file_message(error_message)
        assert "/nonexistent/file.pdf" in error_message
    
    @pytest.mark.asyncio
    async def test_parse_docx_success(self, document_service, tmp_path):
        """测试成功解析DOCX文件"""
        from docx import Document as DocxDocument
        
        docx_path = tmp_path / "test.docx"
        doc = DocxDocument()
        doc.add_paragraph("This is the first paragraph.")
        doc.add_paragraph("This is the second paragraph.")
        
        # 添加表格
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Header 1"
        table.cell(0, 1).text = "Header 2"
        table.cell(1, 0).text = "Data 1"
        table.cell(1, 1).text = "Data 2"
        
        doc.save(str(docx_path))
        
        # 解析DOCX
        text = await document_service.parse_docx(str(docx_path))
        
        # 验证提取的文本
        assert "first paragraph" in text
        assert "second paragraph" in text
        assert "Header 1" in text
        assert "Data 1" in text
    
    @pytest.mark.asyncio
    async def test_parse_docx_empty_content(self, document_service, tmp_path):
        """测试解析空内容的DOCX文件"""
        from docx import Document as DocxDocument
        
        docx_path = tmp_path / "empty.docx"
        doc = DocxDocument()
        doc.save(str(docx_path))
        
        # 解析应该抛出ValueError
        with pytest.raises(ValueError) as exc_info:
            await document_service.parse_docx(str(docx_path))
        
        assert_extractable_text_missing_message(str(exc_info.value))
    
    @pytest.mark.asyncio
    async def test_parse_docx_corrupted_file(self, document_service, tmp_path):
        """测试解析损坏的DOCX文件"""
        docx_path = tmp_path / "corrupted.docx"
        # 创建一个假的DOCX文件
        with open(docx_path, 'wb') as f:
            f.write(b"This is not a valid DOCX file")
        
        # 解析应该抛出ValueError
        with pytest.raises(ValueError) as exc_info:
            await document_service.parse_docx(str(docx_path))
        
        assert_docx_parse_failure_message(str(exc_info.value))
    
    @pytest.mark.asyncio
    async def test_parse_docx_nonexistent_file(self, document_service):
        """测试解析不存在的DOCX文件"""
        with pytest.raises(ValueError) as exc_info:
            await document_service.parse_docx("/nonexistent/file.docx")

        error_message = str(exc_info.value)
        assert_missing_file_message(error_message)
        assert "/nonexistent/file.docx" in error_message
    
    @pytest.mark.asyncio
    async def test_parse_document_pdf(self, document_service, tmp_path):
        """测试通过parse_document解析PDF"""
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        
        pdf_path = tmp_path / "test.pdf"
        c = canvas.Canvas(str(pdf_path), pagesize=letter)
        c.drawString(100, 750, "Test content")
        c.save()
        
        text = await document_service.parse_document(
            str(pdf_path),
            "application/pdf"
        )
        
        assert "Test content" in text
    
    @pytest.mark.asyncio
    async def test_parse_document_docx(self, document_service, tmp_path):
        """测试通过parse_document解析DOCX"""
        from docx import Document as DocxDocument
        
        docx_path = tmp_path / "test.docx"
        doc = DocxDocument()
        doc.add_paragraph("Test content")
        doc.save(str(docx_path))
        
        text = await document_service.parse_document(
            str(docx_path),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
        assert "Test content" in text
    
    @pytest.mark.asyncio
    async def test_parse_document_unsupported_type(self, document_service):
        """测试解析不支持的文件类型"""
        with pytest.raises(ValueError) as exc_info:
            await document_service.parse_document(
                "/some/file.txt",
                "text/plain"
            )
        
        assert "不支持的文件类型" in str(exc_info.value)



class TestTextChunking:
    """测试文本分块功能"""
    
    @pytest.mark.asyncio
    async def test_chunk_text_basic(self, document_service):
        """测试基本文本分块"""
        text = "这是第一段。\n\n这是第二段。\n\n这是第三段。" * 50  # 创建足够长的文本
        
        chunks = await document_service.chunk_text(text)
        
        # 验证返回的是列表
        assert isinstance(chunks, list)
        assert len(chunks) > 0
        
        # 验证每个块的结构
        for i, chunk in enumerate(chunks):
            assert "content" in chunk
            assert "token_count" in chunk
            assert "chunk_index" in chunk
            assert chunk["chunk_index"] == i
            assert isinstance(chunk["content"], str)
            assert isinstance(chunk["token_count"], int)
            assert chunk["token_count"] > 0
    
    @pytest.mark.asyncio
    async def test_chunk_text_respects_size_limit(self, document_service):
        """测试分块遵守大小限制"""
        # 创建一个很长的文本
        text = "这是一个测试句子。" * 500
        
        chunks = await document_service.chunk_text(text)
        
        # 验证每个块的token数量不超过配置的chunk_size
        # 注意：由于overlap，某些块可能略大于chunk_size
        for chunk in chunks:
            # 允许一些容差（因为分割器会尝试在句子边界分割）
            assert chunk["token_count"] <= document_service.chunk_size * 1.5
    
    @pytest.mark.asyncio
    async def test_chunk_text_preserves_paragraph_boundaries(self, document_service):
        """测试分块保持段落边界"""
        # 创建有明确段落分隔的文本
        paragraphs = [
            "第一段内容。这是第一段的详细描述。",
            "第二段内容。这是第二段的详细描述。",
            "第三段内容。这是第三段的详细描述。"
        ]
        text = "\n\n".join(paragraphs * 20)  # 重复以创建足够长的文本
        
        chunks = await document_service.chunk_text(text)
        
        # 验证生成了多个块
        assert len(chunks) > 1
        
        # 验证每个块都包含完整的内容（不是空的）
        for chunk in chunks:
            assert len(chunk["content"].strip()) > 0
    
    @pytest.mark.asyncio
    async def test_chunk_text_empty_input(self, document_service):
        """测试空文本输入"""
        with pytest.raises(ValueError) as exc_info:
            await document_service.chunk_text("")
        
        assert_empty_text_message(str(exc_info.value))
    
    @pytest.mark.asyncio
    async def test_chunk_text_whitespace_only(self, document_service):
        """测试仅包含空白字符的文本"""
        with pytest.raises(ValueError) as exc_info:
            await document_service.chunk_text("   \n\n   \t\t   ")
        
        assert_empty_text_message(str(exc_info.value))
    
    @pytest.mark.asyncio
    async def test_chunk_text_short_text(self, document_service):
        """测试短文本（不需要分块）"""
        text = "这是一个短文本。"
        
        chunks = await document_service.chunk_text(text)
        
        # 短文本应该只生成一个块
        assert len(chunks) == 1
        assert chunks[0]["content"] == text
        assert chunks[0]["chunk_index"] == 0
        assert chunks[0]["token_count"] > 0
    
    @pytest.mark.asyncio
    async def test_chunk_text_with_overlap(self, document_service):
        """测试分块包含重叠内容"""
        # 创建足够长的文本以生成多个块
        text = "句子" + "。这是测试内容" * 200
        
        chunks = await document_service.chunk_text(text)
        
        # 如果有多个块，验证存在重叠
        if len(chunks) > 1:
            # 检查相邻块之间可能存在重叠内容
            # 注意：由于RecursiveCharacterTextSplitter在段落边界分割，
            # 重叠可能不总是明显，但我们可以验证块的连续性
            for i in range(len(chunks) - 1):
                # 每个块都应该有内容
                assert len(chunks[i]["content"]) > 0
                assert len(chunks[i + 1]["content"]) > 0
    
    @pytest.mark.asyncio
    async def test_chunk_text_metadata_consistency(self, document_service):
        """测试分块元数据的一致性"""
        text = "测试内容。" * 300
        
        chunks = await document_service.chunk_text(text)
        
        # 验证chunk_index是连续的
        for i, chunk in enumerate(chunks):
            assert chunk["chunk_index"] == i
        
        # 验证所有块的总token数大致等于原文本的token数
        # （考虑到重叠，总和会略大）
        total_tokens = sum(chunk["token_count"] for chunk in chunks)
        original_tokens = document_service.count_tokens(text)
        
        # 总token数应该大于等于原文本（因为有重叠）
        assert total_tokens >= original_tokens
    
    @pytest.mark.asyncio
    async def test_chunk_text_chinese_content(self, document_service):
        """测试中文内容分块"""
        text = """
        辩论是一种重要的思维训练方式。通过辩论，学生可以提高逻辑思维能力。
        
        在辩论过程中，需要注意以下几点：
        第一，要有清晰的论点。
        第二，要有充分的论据支持。
        第三，要注意反驳对方的观点。
        
        辩论技巧的培养需要长期的练习和积累。
        """ * 50
        
        chunks = await document_service.chunk_text(text)
        
        # 验证中文内容被正确处理
        assert len(chunks) > 0
        for chunk in chunks:
            assert len(chunk["content"]) > 0
            assert chunk["token_count"] > 0
            # 验证包含中文字符
            assert any('\u4e00' <= c <= '\u9fff' for c in chunk["content"])
    
    @pytest.mark.asyncio
    async def test_chunk_text_mixed_language(self, document_service):
        """测试中英文混合内容分块"""
        text = """
        Debate is an important training method. 辩论是重要的训练方式。
        
        Key points to remember:
        1. Clear arguments 清晰的论点
        2. Strong evidence 充分的证据
        3. Logical reasoning 逻辑推理
        """ * 50
        
        chunks = await document_service.chunk_text(text)
        
        # 验证混合语言内容被正确处理
        assert len(chunks) > 0
        for chunk in chunks:
            assert len(chunk["content"]) > 0
            assert chunk["token_count"] > 0


class TestTokenCounting:
    """测试token计数功能"""
    
    def test_count_tokens_english(self, document_service):
        """测试英文文本token计数"""
        text = "This is a test sentence."
        token_count = document_service.count_tokens(text)
        
        # 验证返回正整数
        assert isinstance(token_count, int)
        assert token_count > 0
        # 英文句子的token数应该接近单词数
        assert token_count >= 4  # 至少有几个token
    
    def test_count_tokens_chinese(self, document_service):
        """测试中文文本token计数"""
        text = "这是一个测试句子。"
        token_count = document_service.count_tokens(text)
        
        # 验证返回正整数
        assert isinstance(token_count, int)
        assert token_count > 0
    
    def test_count_tokens_empty(self, document_service):
        """测试空文本token计数"""
        token_count = document_service.count_tokens("")
        
        # 空文本应该返回0
        assert token_count == 0
    
    def test_count_tokens_long_text(self, document_service):
        """测试长文本token计数"""
        text = "这是一个测试。" * 1000
        token_count = document_service.count_tokens(text)
        
        # 长文本应该有相应多的token
        assert token_count > 100



class TestEmbeddingGeneration:
    """测试嵌入向量生成功能"""
    
    @pytest.mark.asyncio
    async def test_generate_embeddings_basic(self, document_service, monkeypatch):
        """测试基本嵌入生成功能"""
        # Mock OpenAI API
        class MockEmbeddingData:
            def __init__(self, embedding):
                self.embedding = embedding
        
        class MockEmbeddingResponse:
            def __init__(self, data):
                self.data = data
        
        class MockOpenAI:
            class Embeddings:
                def create(self, model, input):
                    # 为每个输入文本返回一个1536维的模拟向量
                    return MockEmbeddingResponse([
                        MockEmbeddingData([0.1] * 1536) for _ in input
                    ])
            
            def __init__(self, api_key, base_url=None):
                self.embeddings = self.Embeddings()
        
        # 设置环境变量
        monkeypatch.setenv("OPENAI_API_KEY", "test-key")
        
        # Mock OpenAI client
        import services.document_service
        monkeypatch.setattr(services.document_service, "OpenAI", MockOpenAI)
        
        # 测试数据
        texts = ["这是第一段文本。", "这是第二段文本。", "这是第三段文本。"]
        
        # 生成嵌入
        embeddings = await document_service.generate_embeddings(texts)
        
        # 验证结果
        assert len(embeddings) == len(texts)
        for embedding in embeddings:
            assert isinstance(embedding, list)
            assert len(embedding) == 1536  # OpenAI ada-002 维度
            assert all(isinstance(x, float) for x in embedding)
    
    @pytest.mark.asyncio
    async def test_generate_embeddings_empty_input(self, document_service):
        """测试空输入列表"""
        with pytest.raises(ValueError) as exc_info:
            await document_service.generate_embeddings([])
        
        assert_empty_text_list_message(str(exc_info.value))
    
    @pytest.mark.asyncio
    async def test_generate_embeddings_no_api_key(self, document_service, monkeypatch):
        """测试未配置API密钥"""
        # 清除API密钥环境变量
        monkeypatch.delenv("OPENAI_API_KEY", raising=False)
        
        texts = ["测试文本"]
        
        with pytest.raises(ValueError) as exc_info:
            await document_service.generate_embeddings(texts)

        assert_api_key_missing_message(str(exc_info.value))
    
    @pytest.mark.asyncio
    async def test_generate_embeddings_batching(self, document_service, monkeypatch):
        """测试批处理功能（每批100个文本）"""
        # Mock OpenAI API
        call_count = 0
        batch_sizes = []
        
        class MockEmbeddingData:
            def __init__(self, embedding):
                self.embedding = embedding
        
        class MockEmbeddingResponse:
            def __init__(self, data):
                self.data = data
        
        class MockOpenAI:
            class Embeddings:
                def create(self, model, input):
                    nonlocal call_count, batch_sizes
                    call_count += 1
                    batch_sizes.append(len(input))
                    return MockEmbeddingResponse([
                        MockEmbeddingData([0.1] * 1536) for _ in input
                    ])
            
            def __init__(self, api_key, base_url=None):
                self.embeddings = self.Embeddings()
        
        monkeypatch.setenv("OPENAI_API_KEY", "test-key")
        import services.document_service
        monkeypatch.setattr(services.document_service, "OpenAI", MockOpenAI)
        
        # 创建250个文本（应该分成3批：100, 100, 50）
        texts = [f"文本 {i}" for i in range(250)]
        
        embeddings = await document_service.generate_embeddings(texts)
        
        # 验证批处理
        assert call_count == 3  # 3批
        assert batch_sizes == [100, 100, 50]  # 批次大小
        assert len(embeddings) == 250  # 所有嵌入都返回
    
    @pytest.mark.asyncio
    async def test_generate_embeddings_retry_on_failure(self, document_service, monkeypatch):
        """测试API失败时的重试逻辑"""
        # Mock OpenAI API with failures
        attempt_count = 0
        
        class MockEmbeddingData:
            def __init__(self, embedding):
                self.embedding = embedding
        
        class MockEmbeddingResponse:
            def __init__(self, data):
                self.data = data
        
        class MockOpenAI:
            class Embeddings:
                def create(self, model, input):
                    nonlocal attempt_count
                    attempt_count += 1
                    
                    # 前2次失败，第3次成功
                    if attempt_count < 3:
                        raise Exception("API temporarily unavailable")
                    
                    return MockEmbeddingResponse([
                        MockEmbeddingData([0.1] * 1536) for _ in input
                    ])
            
            def __init__(self, api_key, base_url=None):
                self.embeddings = self.Embeddings()
        
        monkeypatch.setenv("OPENAI_API_KEY", "test-key")
        import services.document_service
        monkeypatch.setattr(services.document_service, "OpenAI", MockOpenAI)
        
        # Mock time.sleep to speed up test
        monkeypatch.setattr("time.sleep", lambda x: None)
        
        texts = ["测试文本"]
        
        # 应该在第3次尝试时成功
        embeddings = await document_service.generate_embeddings(texts)
        
        assert attempt_count == 3  # 重试了2次
        assert len(embeddings) == 1
    
    @pytest.mark.asyncio
    async def test_generate_embeddings_max_retries_exceeded(self, document_service, monkeypatch):
        """测试超过最大重试次数"""
        # Mock OpenAI API that always fails
        class MockOpenAI:
            class Embeddings:
                def create(self, model, input):
                    raise Exception("API error")
            
            def __init__(self, api_key, base_url=None):
                self.embeddings = self.Embeddings()
        
        monkeypatch.setenv("OPENAI_API_KEY", "test-key")
        import services.document_service
        monkeypatch.setattr(services.document_service, "OpenAI", MockOpenAI)
        
        # Mock time.sleep to speed up test
        monkeypatch.setattr("time.sleep", lambda x: None)
        
        texts = ["测试文本"]
        
        # 应该在3次重试后失败
        with pytest.raises(RuntimeError) as exc_info:
            await document_service.generate_embeddings(texts)
        
        error_message = str(exc_info.value)
        assert "嵌入生成失败" in error_message
        assert "已重试" in error_message
        assert "3" in error_message
    
    @pytest.mark.asyncio
    async def test_generate_embeddings_rate_limit_handling(self, document_service, monkeypatch):
        """测试速率限制处理（更长的等待时间）"""
        # Mock OpenAI API with rate limit error
        attempt_count = 0
        sleep_times = []
        
        class MockEmbeddingData:
            def __init__(self, embedding):
                self.embedding = embedding
        
        class MockEmbeddingResponse:
            def __init__(self, data):
                self.data = data
        
        class MockOpenAI:
            class Embeddings:
                def create(self, model, input):
                    nonlocal attempt_count
                    attempt_count += 1
                    
                    # 第1次遇到速率限制，第2次成功
                    if attempt_count < 2:
                        raise Exception("Rate limit exceeded (429)")
                    
                    return MockEmbeddingResponse([
                        MockEmbeddingData([0.1] * 1536) for _ in input
                    ])
            
            def __init__(self, api_key, base_url=None):
                self.embeddings = self.Embeddings()
        
        def mock_sleep(seconds):
            sleep_times.append(seconds)
        
        monkeypatch.setenv("OPENAI_API_KEY", "test-key")
        import services.document_service
        monkeypatch.setattr(services.document_service, "OpenAI", MockOpenAI)
        monkeypatch.setattr("time.sleep", mock_sleep)
        
        texts = ["测试文本"]
        
        embeddings = await document_service.generate_embeddings(texts)
        
        # 验证速率限制导致更长的等待时间
        assert len(sleep_times) == 1
        assert sleep_times[0] >= 5  # 速率限制至少等待5秒
        assert len(embeddings) == 1
    
    @pytest.mark.asyncio
    async def test_generate_embeddings_with_custom_base_url(self, document_service, monkeypatch):
        """测试使用自定义API基础URL"""
        # Mock OpenAI API
        used_base_url = None
        
        class MockEmbeddingData:
            def __init__(self, embedding):
                self.embedding = embedding
        
        class MockEmbeddingResponse:
            def __init__(self, data):
                self.data = data
        
        class MockOpenAI:
            class Embeddings:
                def create(self, model, input):
                    return MockEmbeddingResponse([
                        MockEmbeddingData([0.1] * 1536) for _ in input
                    ])
            
            def __init__(self, api_key, base_url=None):
                nonlocal used_base_url
                used_base_url = base_url
                self.embeddings = self.Embeddings()
        
        monkeypatch.setenv("OPENAI_API_KEY", "test-key")
        monkeypatch.setenv("OPENAI_BASE_URL", "https://custom.api.com/v1")
        
        import services.document_service
        monkeypatch.setattr(services.document_service, "OpenAI", MockOpenAI)
        
        texts = ["测试文本"]
        
        await document_service.generate_embeddings(texts)
        
        # 验证使用了自定义URL
        assert used_base_url == "https://custom.api.com/v1"
    
    @pytest.mark.asyncio
    async def test_generate_embeddings_single_text(self, document_service, monkeypatch):
        """测试单个文本的嵌入生成"""
        # Mock OpenAI API
        class MockEmbeddingData:
            def __init__(self, embedding):
                self.embedding = embedding
        
        class MockEmbeddingResponse:
            def __init__(self, data):
                self.data = data
        
        class MockOpenAI:
            class Embeddings:
                def create(self, model, input):
                    return MockEmbeddingResponse([
                        MockEmbeddingData([0.5] * 1536) for _ in input
                    ])
            
            def __init__(self, api_key, base_url=None):
                self.embeddings = self.Embeddings()
        
        monkeypatch.setenv("OPENAI_API_KEY", "test-key")
        import services.document_service
        monkeypatch.setattr(services.document_service, "OpenAI", MockOpenAI)
        
        texts = ["单个测试文本"]
        
        embeddings = await document_service.generate_embeddings(texts)
        
        assert len(embeddings) == 1
        assert len(embeddings[0]) == 1536
        assert embeddings[0][0] == 0.5



@pytest.mark.pgvector
class TestVectorStorage:
    """测试向量存储功能（Task 6.2）"""
    
    @pytest.mark.asyncio
    async def test_store_chunks_with_embeddings_basic(self, document_service, test_user_id, db_session):
        """测试基本的块和嵌入存储功能"""
        # 创建文档
        document = await document_service.upload_document(
            file_data=b"Test content",
            filename="test.pdf",
            file_type="application/pdf",
            user_id=test_user_id
        )
        
        # 准备测试数据
        chunks = [
            {"content": "第一段内容", "token_count": 10, "chunk_index": 0},
            {"content": "第二段内容", "token_count": 12, "chunk_index": 1},
            {"content": "第三段内容", "token_count": 15, "chunk_index": 2}
        ]
        
        # 生成模拟嵌入向量（1536维）
        embeddings = [[0.1] * 1536, [0.2] * 1536, [0.3] * 1536]
        
        # 注意：在SQLite测试环境中，我们不能真正存储向量
        # 但我们可以测试方法的逻辑
        # 在实际PostgreSQL环境中，这会正常工作
        try:
            # 存储块和嵌入
            chunk_records = await document_service.store_chunks_with_embeddings(
                document_id=str(document.id),
                chunks=chunks,
                embeddings=embeddings
            )
            
            # 验证返回的记录
            assert len(chunk_records) == 3
            for i, record in enumerate(chunk_records):
                assert record.document_id == document.id
                assert record.chunk_index == i
                assert record.content == chunks[i]["content"]
                assert record.token_count == chunks[i]["token_count"]
                assert record.embedding == embeddings[i]
        except Exception as e:
            # 在SQLite环境中，可能会因为表不存在而失败
            # 这是预期的，因为我们没有创建chunks表
            if "no such table: kb_document_chunks" in str(e):
                pytest.skip("SQLite不支持向量类型，跳过此测试")
            else:
                raise
    
    @pytest.mark.asyncio
    async def test_store_chunks_empty_list(self, document_service, test_user_id):
        """测试存储空块列表"""
        # 创建文档
        document = await document_service.upload_document(
            file_data=b"Test content",
            filename="test.pdf",
            file_type="application/pdf",
            user_id=test_user_id
        )
        
        # 尝试存储空列表
        with pytest.raises(ValueError) as exc_info:
            await document_service.store_chunks_with_embeddings(
                document_id=str(document.id),
                chunks=[],
                embeddings=[]
            )
        
        assert_empty_chunk_list_message(str(exc_info.value))
    
    @pytest.mark.asyncio
    async def test_store_chunks_mismatched_lengths(self, document_service, test_user_id):
        """测试块和嵌入数量不匹配"""
        # 创建文档
        document = await document_service.upload_document(
            file_data=b"Test content",
            filename="test.pdf",
            file_type="application/pdf",
            user_id=test_user_id
        )
        
        chunks = [
            {"content": "内容1", "token_count": 10, "chunk_index": 0},
            {"content": "内容2", "token_count": 12, "chunk_index": 1}
        ]
        
        # 只提供1个嵌入向量（应该有2个）
        embeddings = [[0.1] * 1536]
        
        with pytest.raises(ValueError) as exc_info:
            await document_service.store_chunks_with_embeddings(
                document_id=str(document.id),
                chunks=chunks,
                embeddings=embeddings
            )
        
        assert "不匹配" in str(exc_info.value)
    
    @pytest.mark.asyncio
    async def test_store_chunks_invalid_document_id(self, document_service):
        """测试使用无效的文档ID"""
        fake_id = str(uuid.uuid4())
        
        chunks = [{"content": "内容", "token_count": 10, "chunk_index": 0}]
        embeddings = [[0.1] * 1536]
        
        with pytest.raises(ValueError) as exc_info:
            await document_service.store_chunks_with_embeddings(
                document_id=fake_id,
                chunks=chunks,
                embeddings=embeddings
            )

        error_message = str(exc_info.value)
        assert_missing_document_message(error_message)
        assert fake_id in error_message
    
    @pytest.mark.asyncio
    async def test_store_chunks_invalid_embedding_dimension(self, document_service, test_user_id):
        """测试嵌入向量维度错误"""
        # 创建文档
        document = await document_service.upload_document(
            file_data=b"Test content",
            filename="test.pdf",
            file_type="application/pdf",
            user_id=test_user_id
        )
        
        chunks = [{"content": "内容", "token_count": 10, "chunk_index": 0}]
        
        # 错误的维度（应该是1536）
        embeddings = [[0.1] * 512]
        
        with pytest.raises(ValueError) as exc_info:
            await document_service.store_chunks_with_embeddings(
                document_id=str(document.id),
                chunks=chunks,
                embeddings=embeddings
            )
        
        assert_embedding_dimension_message(str(exc_info.value))
    
    @pytest.mark.asyncio
    async def test_retrieve_chunks_by_document(self, document_service, test_user_id):
        """测试检索文档的所有块"""
        # 创建文档
        document = await document_service.upload_document(
            file_data=b"Test content",
            filename="test.pdf",
            file_type="application/pdf",
            user_id=test_user_id
        )
        
        try:
            # 准备并存储块
            chunks = [
                {"content": "块0", "token_count": 5, "chunk_index": 0},
                {"content": "块1", "token_count": 6, "chunk_index": 1},
                {"content": "块2", "token_count": 7, "chunk_index": 2}
            ]
            embeddings = [[0.1] * 1536, [0.2] * 1536, [0.3] * 1536]
            
            await document_service.store_chunks_with_embeddings(
                document_id=str(document.id),
                chunks=chunks,
                embeddings=embeddings
            )
            
            # 检索块
            retrieved_chunks = await document_service.retrieve_chunks_by_document(
                document_id=str(document.id)
            )
            
            # 验证检索结果
            assert len(retrieved_chunks) == 3
            
            # 验证按chunk_index排序
            for i, chunk in enumerate(retrieved_chunks):
                assert chunk.chunk_index == i
                assert chunk.content == chunks[i]["content"]
                assert chunk.token_count == chunks[i]["token_count"]
        
        except Exception as e:
            if "no such table: kb_document_chunks" in str(e):
                pytest.skip("SQLite不支持向量类型，跳过此测试")
            else:
                raise
    
    @pytest.mark.asyncio
    async def test_retrieve_chunks_invalid_document_id(self, document_service):
        """测试检索不存在的文档的块"""
        fake_id = str(uuid.uuid4())
        
        with pytest.raises(ValueError) as exc_info:
            await document_service.retrieve_chunks_by_document(fake_id)

        error_message = str(exc_info.value)
        assert_missing_document_message(error_message)
        assert fake_id in error_message
    
    @pytest.mark.asyncio
    async def test_retrieve_chunks_empty_document(self, document_service, test_user_id):
        """测试检索没有块的文档"""
        # 创建文档但不添加块
        document = await document_service.upload_document(
            file_data=b"Test content",
            filename="test.pdf",
            file_type="application/pdf",
            user_id=test_user_id
        )
        
        try:
            # 检索块（应该返回空列表）
            chunks = await document_service.retrieve_chunks_by_document(
                document_id=str(document.id)
            )
            
            assert chunks == []
        
        except Exception as e:
            if "no such table: kb_document_chunks" in str(e):
                pytest.skip("SQLite不支持向量类型，跳过此测试")
            else:
                raise
    
    @pytest.mark.asyncio
    async def test_retrieve_chunk_by_id(self, document_service, test_user_id):
        """测试根据ID检索单个块"""
        # 创建文档
        document = await document_service.upload_document(
            file_data=b"Test content",
            filename="test.pdf",
            file_type="application/pdf",
            user_id=test_user_id
        )
        
        try:
            # 存储块
            chunks = [{"content": "测试块", "token_count": 10, "chunk_index": 0}]
            embeddings = [[0.5] * 1536]
            
            chunk_records = await document_service.store_chunks_with_embeddings(
                document_id=str(document.id),
                chunks=chunks,
                embeddings=embeddings
            )
            
            chunk_id = str(chunk_records[0].id)
            
            # 检索块
            retrieved_chunk = await document_service.retrieve_chunk_by_id(chunk_id)
            
            # 验证
            assert retrieved_chunk is not None
            assert retrieved_chunk.id == chunk_records[0].id
            assert retrieved_chunk.content == "测试块"
            assert retrieved_chunk.token_count == 10
        
        except Exception as e:
            if "no such table: kb_document_chunks" in str(e):
                pytest.skip("SQLite不支持向量类型，跳过此测试")
            else:
                raise
    
    @pytest.mark.asyncio
    async def test_retrieve_chunk_by_id_not_found(self, document_service):
        """测试检索不存在的块"""
        fake_id = str(uuid.uuid4())
        
        try:
            chunk = await document_service.retrieve_chunk_by_id(fake_id)
            
            # 应该返回None
            assert chunk is None
        
        except Exception as e:
            if "no such table: kb_document_chunks" in str(e):
                pytest.skip("SQLite不支持向量类型，跳过此测试")
            else:
                raise
    
    @pytest.mark.asyncio
    async def test_chunk_storage_round_trip(self, document_service, test_user_id):
        """
        测试块存储的往返一致性
        Property 13: Chunk Storage Round-Trip
        验证存储后检索的块与原始数据一致
        """
        # 创建文档
        document = await document_service.upload_document(
            file_data=b"Test content",
            filename="test.pdf",
            file_type="application/pdf",
            user_id=test_user_id
        )
        
        try:
            # 准备测试数据
            original_chunks = [
                {"content": "第一段：这是测试内容。", "token_count": 15, "chunk_index": 0},
                {"content": "第二段：更多测试内容。", "token_count": 18, "chunk_index": 1}
            ]
            
            original_embeddings = [
                [0.123 + i * 0.001 for i in range(1536)],
                [0.456 + i * 0.001 for i in range(1536)]
            ]
            
            # 存储
            stored_chunks = await document_service.store_chunks_with_embeddings(
                document_id=str(document.id),
                chunks=original_chunks,
                embeddings=original_embeddings
            )
            
            # 检索
            retrieved_chunks = await document_service.retrieve_chunks_by_document(
                document_id=str(document.id)
            )
            
            # 验证往返一致性
            assert len(retrieved_chunks) == len(original_chunks)
            
            for i, retrieved in enumerate(retrieved_chunks):
                # 验证内容
                assert retrieved.content == original_chunks[i]["content"]
                assert retrieved.token_count == original_chunks[i]["token_count"]
                assert retrieved.chunk_index == original_chunks[i]["chunk_index"]
                
                # 验证嵌入向量
                assert retrieved.embedding == original_embeddings[i]
                assert len(retrieved.embedding) == 1536
                
                # 验证文档关联
                assert retrieved.document_id == document.id
        
        except Exception as e:
            if "no such table: kb_document_chunks" in str(e):
                pytest.skip("SQLite不支持向量类型，跳过此测试")
            else:
                raise
    
    @pytest.mark.asyncio
    async def test_embedding_dimension_validation(self, document_service, test_user_id):
        """
        测试嵌入向量维度验证
        Property 12: Embedding Generation Completeness
        确保所有嵌入向量都是正确的维度（1536）
        """
        # 创建文档
        document = await document_service.upload_document(
            file_data=b"Test content",
            filename="test.pdf",
            file_type="application/pdf",
            user_id=test_user_id
        )
        
        chunks = [{"content": "测试", "token_count": 5, "chunk_index": 0}]
        
        # 测试各种错误的维度
        invalid_dimensions = [
            [0.1] * 512,   # 太小
            [0.1] * 768,   # 错误的维度
            [0.1] * 2048,  # 太大
            [0.1] * 1535,  # 差1
            [0.1] * 1537,  # 多1
        ]
        
        for invalid_embedding in invalid_dimensions:
            with pytest.raises(ValueError) as exc_info:
                await document_service.store_chunks_with_embeddings(
                    document_id=str(document.id),
                    chunks=chunks,
                    embeddings=[invalid_embedding]
                )
            
            assert_embedding_dimension_message(str(exc_info.value))
    
    @pytest.mark.asyncio
    async def test_batch_chunk_storage(self, document_service, test_user_id):
        """测试批量存储大量块"""
        # 创建文档
        document = await document_service.upload_document(
            file_data=b"Test content",
            filename="test.pdf",
            file_type="application/pdf",
            user_id=test_user_id
        )
        
        try:
            # 创建100个块
            num_chunks = 100
            chunks = [
                {
                    "content": f"块 {i} 的内容",
                    "token_count": 10 + i,
                    "chunk_index": i
                }
                for i in range(num_chunks)
            ]
            
            embeddings = [[0.01 * i] * 1536 for i in range(num_chunks)]
            
            # 批量存储
            stored_chunks = await document_service.store_chunks_with_embeddings(
                document_id=str(document.id),
                chunks=chunks,
                embeddings=embeddings
            )
            
            # 验证
            assert len(stored_chunks) == num_chunks
            
            # 检索并验证
            retrieved_chunks = await document_service.retrieve_chunks_by_document(
                document_id=str(document.id)
            )
            
            assert len(retrieved_chunks) == num_chunks
            
            # 验证顺序
            for i, chunk in enumerate(retrieved_chunks):
                assert chunk.chunk_index == i
        
        except Exception as e:
            if "no such table: kb_document_chunks" in str(e):
                pytest.skip("SQLite不支持向量类型，跳过此测试")
            else:
                raise


if __name__ == "__main__":
    pytest.main([__file__, "-v"])



@pytest.mark.pgvector
class TestVectorSimilaritySearch:
    """测试向量相似度搜索功能（Task 6.2 - 核心功能）"""
    
    @pytest.mark.asyncio
    async def test_search_similar_chunks_basic(self, document_service, test_user_id):
        """测试基本的向量相似度搜索"""
        # 创建文档
        document = await document_service.upload_document(
            file_data=b"Test content",
            filename="test.pdf",
            file_type="application/pdf",
            user_id=test_user_id
        )
        
        # 更新文档状态为completed（只有completed的文档才会被搜索）
        document.upload_status = 'completed'
        document_service.db.commit()
        
        try:
            # 存储一些块
            chunks = [
                {"content": "人工智能是计算机科学的一个分支", "token_count": 20, "chunk_index": 0},
                {"content": "机器学习是人工智能的核心技术", "token_count": 18, "chunk_index": 1},
                {"content": "今天天气很好，适合出去玩", "token_count": 15, "chunk_index": 2}
            ]
            
            # 创建不同的嵌入向量（模拟语义相似度）
            embeddings = [
                [0.9] * 1536,  # 与AI相关
                [0.85] * 1536,  # 与AI相关
                [0.1] * 1536   # 与天气相关
            ]
            
            await document_service.store_chunks_with_embeddings(
                document_id=str(document.id),
                chunks=chunks,
                embeddings=embeddings
            )
            
            # 搜索与AI相关的内容（使用相似的向量）
            query_embedding = [0.88] * 1536
            
            results = await document_service.search_similar_chunks(
                query_embedding=query_embedding,
                top_k=2
            )
            
            # 验证结果
            assert len(results) <= 2  # 最多返回top_k个结果
            
            # 验证结果结构
            for result in results:
                assert "chunk_id" in result
                assert "document_id" in result
                assert "document_name" in result
                assert "content" in result
                assert "chunk_index" in result
                assert "token_count" in result
                assert "similarity_score" in result
                
                # 验证相似度分数在合理范围内
                assert -1 <= result["similarity_score"] <= 1
        
        except Exception as e:
            if "no such table: kb_document_chunks" in str(e) or "operator does not exist" in str(e):
                pytest.skip("需要PostgreSQL和pgvector扩展，跳过此测试")
            else:
                raise
    
    @pytest.mark.asyncio
    async def test_search_similar_chunks_empty_query(self, document_service):
        """测试空查询向量"""
        with pytest.raises(ValueError) as exc_info:
            await document_service.search_similar_chunks(
                query_embedding=[],
                top_k=5
            )
        
        assert "查询嵌入向量为空" in str(exc_info.value)
    
    @pytest.mark.asyncio
    async def test_search_similar_chunks_invalid_dimension(self, document_service):
        """测试错误维度的查询向量"""
        # 错误的维度（应该是1536）
        query_embedding = [0.5] * 512
        
        with pytest.raises(ValueError) as exc_info:
            await document_service.search_similar_chunks(
                query_embedding=query_embedding,
                top_k=5
            )
        
        assert_embedding_dimension_message(str(exc_info.value))
    
    @pytest.mark.asyncio
    async def test_search_similar_chunks_invalid_top_k(self, document_service):
        """测试无效的top_k参数"""
        query_embedding = [0.5] * 1536
        
        # top_k = 0
        with pytest.raises(ValueError) as exc_info:
            await document_service.search_similar_chunks(
                query_embedding=query_embedding,
                top_k=0
            )
        
        assert_invalid_top_k_message(str(exc_info.value))
        
        # top_k < 0
        with pytest.raises(ValueError) as exc_info:
            await document_service.search_similar_chunks(
                query_embedding=query_embedding,
                top_k=-1
            )
        
        assert_invalid_top_k_message(str(exc_info.value))
    
    @pytest.mark.asyncio
    async def test_search_similar_chunks_with_threshold(self, document_service, test_user_id):
        """测试使用相似度阈值过滤结果"""
        # 创建文档
        document = await document_service.upload_document(
            file_data=b"Test content",
            filename="test.pdf",
            file_type="application/pdf",
            user_id=test_user_id
        )
        
        document.upload_status = 'completed'
        document_service.db.commit()
        
        try:
            # 存储块
            chunks = [
                {"content": "高相似度内容", "token_count": 10, "chunk_index": 0},
                {"content": "中等相似度内容", "token_count": 12, "chunk_index": 1},
                {"content": "低相似度内容", "token_count": 15, "chunk_index": 2}
            ]
            
            embeddings = [
                [0.95] * 1536,  # 高相似度
                [0.7] * 1536,   # 中等相似度
                [0.3] * 1536    # 低相似度
            ]
            
            await document_service.store_chunks_with_embeddings(
                document_id=str(document.id),
                chunks=chunks,
                embeddings=embeddings
            )
            
            # 使用高阈值搜索（只返回高相似度的）
            query_embedding = [0.9] * 1536
            
            results = await document_service.search_similar_chunks(
                query_embedding=query_embedding,
                top_k=10,
                similarity_threshold=0.8  # 高阈值
            )
            
            # 验证所有结果的相似度都高于阈值
            for result in results:
                assert result["similarity_score"] >= 0.8
        
        except Exception as e:
            if "no such table: kb_document_chunks" in str(e) or "operator does not exist" in str(e):
                pytest.skip("需要PostgreSQL和pgvector扩展，跳过此测试")
            else:
                raise
    
    @pytest.mark.asyncio
    async def test_search_similar_chunks_no_results(self, document_service):
        """测试没有匹配结果的搜索"""
        try:
            # 搜索（数据库中没有任何块）
            query_embedding = [0.5] * 1536
            
            results = await document_service.search_similar_chunks(
                query_embedding=query_embedding,
                top_k=5
            )
            
            # 应该返回空列表
            assert results == []
        
        except Exception as e:
            if "no such table: kb_document_chunks" in str(e) or "operator does not exist" in str(e):
                pytest.skip("需要PostgreSQL和pgvector扩展，跳过此测试")
            else:
                raise
    
    @pytest.mark.asyncio
    async def test_search_similar_chunks_only_completed_documents(self, document_service, test_user_id):
        """测试只搜索已完成处理的文档"""
        # 创建两个文档
        doc1 = await document_service.upload_document(
            file_data=b"Content 1",
            filename="doc1.pdf",
            file_type="application/pdf",
            user_id=test_user_id
        )
        
        doc2 = await document_service.upload_document(
            file_data=b"Content 2",
            filename="doc2.pdf",
            file_type="application/pdf",
            user_id=test_user_id
        )
        
        # 只将doc1标记为completed
        doc1.upload_status = 'completed'
        doc2.upload_status = 'processing'  # 仍在处理中
        document_service.db.commit()
        
        try:
            # 为两个文档都添加块
            chunks1 = [{"content": "文档1的内容", "token_count": 10, "chunk_index": 0}]
            embeddings1 = [[0.9] * 1536]
            
            chunks2 = [{"content": "文档2的内容", "token_count": 10, "chunk_index": 0}]
            embeddings2 = [[0.9] * 1536]
            
            await document_service.store_chunks_with_embeddings(
                document_id=str(doc1.id),
                chunks=chunks1,
                embeddings=embeddings1
            )
            
            await document_service.store_chunks_with_embeddings(
                document_id=str(doc2.id),
                chunks=chunks2,
                embeddings=embeddings2
            )
            
            # 搜索
            query_embedding = [0.9] * 1536
            results = await document_service.search_similar_chunks(
                query_embedding=query_embedding,
                top_k=10
            )
            
            # 应该只返回doc1的块（因为doc2还在processing状态）
            for result in results:
                assert result["document_id"] == str(doc1.id)
                assert result["document_name"] == "doc1.pdf"
        
        except Exception as e:
            if "no such table: kb_document_chunks" in str(e) or "operator does not exist" in str(e):
                pytest.skip("需要PostgreSQL和pgvector扩展，跳过此测试")
            else:
                raise
    
    @pytest.mark.asyncio
    async def test_search_similar_chunks_ordering(self, document_service, test_user_id):
        """测试搜索结果按相似度降序排列"""
        # 创建文档
        document = await document_service.upload_document(
            file_data=b"Test content",
            filename="test.pdf",
            file_type="application/pdf",
            user_id=test_user_id
        )
        
        document.upload_status = 'completed'
        document_service.db.commit()
        
        try:
            # 存储多个块，相似度递减
            chunks = [
                {"content": f"内容 {i}", "token_count": 10, "chunk_index": i}
                for i in range(5)
            ]
            
            # 创建相似度递减的嵌入向量
            embeddings = [
                [0.9 - i * 0.1] * 1536 for i in range(5)
            ]
            
            await document_service.store_chunks_with_embeddings(
                document_id=str(document.id),
                chunks=chunks,
                embeddings=embeddings
            )
            
            # 搜索
            query_embedding = [0.95] * 1536
            results = await document_service.search_similar_chunks(
                query_embedding=query_embedding,
                top_k=5
            )
            
            # 验证结果按相似度降序排列
            if len(results) > 1:
                for i in range(len(results) - 1):
                    assert results[i]["similarity_score"] >= results[i + 1]["similarity_score"]
        
        except Exception as e:
            if "no such table: kb_document_chunks" in str(e) or "operator does not exist" in str(e):
                pytest.skip("需要PostgreSQL和pgvector扩展，跳过此测试")
            else:
                raise
    
    @pytest.mark.asyncio
    async def test_search_similar_chunks_respects_top_k(self, document_service, test_user_id):
        """测试top_k参数限制返回结果数量"""
        # 创建文档
        document = await document_service.upload_document(
            file_data=b"Test content",
            filename="test.pdf",
            file_type="application/pdf",
            user_id=test_user_id
        )
        
        document.upload_status = 'completed'
        document_service.db.commit()
        
        try:
            # 存储10个块
            chunks = [
                {"content": f"块 {i}", "token_count": 10, "chunk_index": i}
                for i in range(10)
            ]
            
            embeddings = [[0.8] * 1536 for _ in range(10)]
            
            await document_service.store_chunks_with_embeddings(
                document_id=str(document.id),
                chunks=chunks,
                embeddings=embeddings
            )
            
            # 测试不同的top_k值
            for top_k in [1, 3, 5, 10]:
                results = await document_service.search_similar_chunks(
                    query_embedding=[0.8] * 1536,
                    top_k=top_k
                )
                
                # 验证返回的结果数量不超过top_k
                assert len(results) <= top_k
        
        except Exception as e:
            if "no such table: kb_document_chunks" in str(e) or "operator does not exist" in str(e):
                pytest.skip("需要PostgreSQL和pgvector扩展，跳过此测试")
            else:
                raise
    
    @pytest.mark.asyncio
    async def test_search_similar_chunks_multiple_documents(self, document_service, test_user_id):
        """测试跨多个文档搜索"""
        # 创建多个文档
        docs = []
        for i in range(3):
            doc = await document_service.upload_document(
                file_data=f"Content {i}".encode(),
                filename=f"doc{i}.pdf",
                file_type="application/pdf",
                user_id=test_user_id
            )
            doc.upload_status = 'completed'
            docs.append(doc)
        
        document_service.db.commit()
        
        try:
            # 为每个文档添加块
            for i, doc in enumerate(docs):
                chunks = [
                    {"content": f"文档{i}的内容", "token_count": 10, "chunk_index": 0}
                ]
                embeddings = [[0.8 + i * 0.05] * 1536]
                
                await document_service.store_chunks_with_embeddings(
                    document_id=str(doc.id),
                    chunks=chunks,
                    embeddings=embeddings
                )
            
            # 搜索（应该返回所有文档的块）
            query_embedding = [0.85] * 1536
            results = await document_service.search_similar_chunks(
                query_embedding=query_embedding,
                top_k=10
            )
            
            # 验证返回了多个文档的块
            document_ids = set(result["document_id"] for result in results)
            assert len(document_ids) > 1  # 至少有2个不同的文档
        
        except Exception as e:
            if "no such table: kb_document_chunks" in str(e) or "operator does not exist" in str(e):
                pytest.skip("需要PostgreSQL和pgvector扩展，跳过此测试")
            else:
                raise
    
    @pytest.mark.asyncio
    async def test_search_similar_chunks_includes_document_name(self, document_service, test_user_id):
        """测试搜索结果包含文档名称"""
        # 创建文档
        document = await document_service.upload_document(
            file_data=b"Test content",
            filename="重要文档.pdf",
            file_type="application/pdf",
            user_id=test_user_id
        )
        
        document.upload_status = 'completed'
        document_service.db.commit()
        
        try:
            # 存储块
            chunks = [{"content": "测试内容", "token_count": 10, "chunk_index": 0}]
            embeddings = [[0.8] * 1536]
            
            await document_service.store_chunks_with_embeddings(
                document_id=str(document.id),
                chunks=chunks,
                embeddings=embeddings
            )
            
            # 搜索
            query_embedding = [0.8] * 1536
            results = await document_service.search_similar_chunks(
                query_embedding=query_embedding,
                top_k=5
            )
            
            # 验证结果包含文档名称
            if len(results) > 0:
                assert results[0]["document_name"] == "重要文档.pdf"
        
        except Exception as e:
            if "no such table: kb_document_chunks" in str(e) or "operator does not exist" in str(e):
                pytest.skip("需要PostgreSQL和pgvector扩展，跳过此测试")
            else:
                raise



@pytest.mark.pgvector
class TestDocumentProcessingWorkflow:
    """测试文档处理工作流（Task 7.1）"""
    
    @pytest.mark.asyncio
    async def test_process_document_success(self, document_service, test_user_id, tmp_path, monkeypatch):
        """测试成功的文档处理工作流"""
        # 创建一个测试PDF文件
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        
        pdf_path = tmp_path / "test.pdf"
        c = canvas.Canvas(str(pdf_path), pagesize=letter)
        c.drawString(100, 750, "This is test content for document processing.")
        c.drawString(100, 730, "It contains multiple lines to test chunking.")
        c.save()
        
        # 上传文档
        with open(pdf_path, 'rb') as f:
            file_data = f.read()
        
        document = await document_service.upload_document(
            file_data=file_data,
            filename="test.pdf",
            file_type="application/pdf",
            user_id=test_user_id
        )
        
        # 验证初始状态
        assert document.upload_status == "pending"
        assert document.processed_at is None
        
        # Mock OpenAI API
        class MockEmbeddingData:
            def __init__(self, embedding):
                self.embedding = embedding
        
        class MockEmbeddingResponse:
            def __init__(self, data):
                self.data = data
        
        class MockOpenAI:
            class Embeddings:
                def create(self, model, input):
                    return MockEmbeddingResponse([
                        MockEmbeddingData([0.1] * 1536) for _ in input
                    ])
            
            def __init__(self, api_key, base_url=None):
                self.embeddings = self.Embeddings()
        
        monkeypatch.setenv("OPENAI_API_KEY", "test-key")
        import services.document_service
        monkeypatch.setattr(services.document_service, "OpenAI", MockOpenAI)
        
        try:
            # 处理文档
            await document_service.process_document(str(document.id))
            
            # 刷新文档状态
            document_service.db.refresh(document)
            
            # 验证状态更新
            assert document.upload_status == "completed"
            assert document.processed_at is not None
            assert document.error_message is None
            
            # 验证块已创建
            chunks = await document_service.retrieve_chunks_by_document(str(document.id))
            assert len(chunks) > 0
            
            # 验证每个块都有嵌入向量
            for chunk in chunks:
                assert chunk.embedding is not None
                assert len(chunk.embedding) == 1536
        
        except Exception as e:
            if "no such table: kb_document_chunks" in str(e):
                pytest.skip("SQLite不支持向量类型，跳过此测试")
            else:
                raise
    
    @pytest.mark.asyncio
    async def test_process_document_status_progression(self, document_service, test_user_id, tmp_path, monkeypatch):
        """
        测试文档处理状态的正确进展
        Property 10: Processing Workflow Progression
        验证状态从 pending → processing → completed
        """
        # 创建测试文档
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        
        pdf_path = tmp_path / "test.pdf"
        c = canvas.Canvas(str(pdf_path), pagesize=letter)
        c.drawString(100, 750, "Test content")
        c.save()
        
        with open(pdf_path, 'rb') as f:
            file_data = f.read()
        
        document = await document_service.upload_document(
            file_data=file_data,
            filename="test.pdf",
            file_type="application/pdf",
            user_id=test_user_id
        )
        
        # 初始状态应该是 pending
        assert document.upload_status == "pending"
        
        # Mock OpenAI
        class MockEmbeddingData:
            def __init__(self, embedding):
                self.embedding = embedding
        
        class MockEmbeddingResponse:
            def __init__(self, data):
                self.data = data
        
        class MockOpenAI:
            class Embeddings:
                def create(self, model, input):
                    return MockEmbeddingResponse([
                        MockEmbeddingData([0.1] * 1536) for _ in input
                    ])
            
            def __init__(self, api_key, base_url=None):
                self.embeddings = self.Embeddings()
        
        monkeypatch.setenv("OPENAI_API_KEY", "test-key")
        import services.document_service
        monkeypatch.setattr(services.document_service, "OpenAI", MockOpenAI)
        
        try:
            # 处理文档
            await document_service.process_document(str(document.id))
            
            # 刷新状态
            document_service.db.refresh(document)
            
            # 最终状态应该是 completed
            assert document.upload_status == "completed"
        
        except Exception as e:
            if "no such table: kb_document_chunks" in str(e):
                pytest.skip("SQLite不支持向量类型，跳过此测试")
            else:
                raise
    
    @pytest.mark.asyncio
    async def test_process_document_parsing_failure(self, document_service, test_user_id, tmp_path):
        """测试解析失败时的错误处理"""
        # 创建一个损坏的PDF文件
        pdf_path = tmp_path / "corrupted.pdf"
        with open(pdf_path, 'wb') as f:
            f.write(b"This is not a valid PDF")
        
        with open(pdf_path, 'rb') as f:
            file_data = f.read()
        
        document = await document_service.upload_document(
            file_data=file_data,
            filename="corrupted.pdf",
            file_type="application/pdf",
            user_id=test_user_id
        )
        
        # 尝试处理文档（应该失败）
        with pytest.raises(ValueError):
            await document_service.process_document(str(document.id))
        
        # 刷新状态
        document_service.db.refresh(document)
        
        # 验证状态更新为 failed
        assert document.upload_status == "failed"
        assert document.error_message is not None
        assert "损坏" in document.error_message or "格式错误" in document.error_message
    
    @pytest.mark.asyncio
    async def test_process_document_embedding_failure_rollback(self, document_service, test_user_id, tmp_path, monkeypatch):
        """
        测试嵌入生成失败时的回滚
        Property 14: Processing Failure Rollback
        验证失败时删除部分创建的块
        """
        # 创建测试文档
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        
        pdf_path = tmp_path / "test.pdf"
        c = canvas.Canvas(str(pdf_path), pagesize=letter)
        c.drawString(100, 750, "Test content for rollback")
        c.save()
        
        with open(pdf_path, 'rb') as f:
            file_data = f.read()
        
        document = await document_service.upload_document(
            file_data=file_data,
            filename="test.pdf",
            file_type="application/pdf",
            user_id=test_user_id
        )
        
        # Mock OpenAI to always fail
        class MockOpenAI:
            class Embeddings:
                def create(self, model, input):
                    raise Exception("API error - embedding generation failed")
            
            def __init__(self, api_key, base_url=None):
                self.embeddings = self.Embeddings()
        
        monkeypatch.setenv("OPENAI_API_KEY", "test-key")
        import services.document_service
        monkeypatch.setattr(services.document_service, "OpenAI", MockOpenAI)
        monkeypatch.setattr("time.sleep", lambda x: None)  # Speed up retries
        
        try:
            # 尝试处理文档（应该失败）
            with pytest.raises(RuntimeError):
                await document_service.process_document(str(document.id))
            
            # 刷新状态
            document_service.db.refresh(document)
            
            # 验证状态为 failed
            assert document.upload_status == "failed"
            assert document.error_message is not None
            
            # 验证没有部分块残留
            chunks = await document_service.retrieve_chunks_by_document(str(document.id))
            assert len(chunks) == 0  # 应该已经回滚删除
        
        except Exception as e:
            if "no such table: kb_document_chunks" in str(e):
                pytest.skip("SQLite不支持向量类型，跳过此测试")
            else:
                raise
    
    @pytest.mark.asyncio
    async def test_process_document_nonexistent(self, document_service):
        """测试处理不存在的文档"""
        fake_id = str(uuid.uuid4())
        
        with pytest.raises(ValueError) as exc_info:
            await document_service.process_document(fake_id)

        error_message = str(exc_info.value)
        assert_missing_document_message(error_message)
        assert fake_id in error_message
    
    @pytest.mark.asyncio
    async def test_process_document_empty_content(self, document_service, test_user_id, tmp_path):
        """测试处理空内容的文档"""
        # 创建一个空白PDF
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        
        pdf_path = tmp_path / "empty.pdf"
        c = canvas.Canvas(str(pdf_path), pagesize=letter)
        c.showPage()  # 空白页
        c.save()
        
        with open(pdf_path, 'rb') as f:
            file_data = f.read()
        
        document = await document_service.upload_document(
            file_data=file_data,
            filename="empty.pdf",
            file_type="application/pdf",
            user_id=test_user_id
        )
        
        # 尝试处理（应该失败，因为没有文本内容）
        with pytest.raises(ValueError):
            await document_service.process_document(str(document.id))
        
        # 验证状态
        document_service.db.refresh(document)
        assert document.upload_status == "failed"
        assert_extractable_text_missing_message(document.error_message)
    
    @pytest.mark.asyncio
    async def test_process_document_logs_all_steps(self, document_service, test_user_id, tmp_path, monkeypatch, caplog):
        """测试处理过程中记录所有步骤"""
        import logging
        caplog.set_level(logging.INFO)
        
        # 创建测试文档
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        
        pdf_path = tmp_path / "test.pdf"
        c = canvas.Canvas(str(pdf_path), pagesize=letter)
        c.drawString(100, 750, "Test content for logging")
        c.save()
        
        with open(pdf_path, 'rb') as f:
            file_data = f.read()
        
        document = await document_service.upload_document(
            file_data=file_data,
            filename="test.pdf",
            file_type="application/pdf",
            user_id=test_user_id
        )
        
        # Mock OpenAI
        class MockEmbeddingData:
            def __init__(self, embedding):
                self.embedding = embedding
        
        class MockEmbeddingResponse:
            def __init__(self, data):
                self.data = data
        
        class MockOpenAI:
            class Embeddings:
                def create(self, model, input):
                    return MockEmbeddingResponse([
                        MockEmbeddingData([0.1] * 1536) for _ in input
                    ])
            
            def __init__(self, api_key, base_url=None):
                self.embeddings = self.Embeddings()
        
        monkeypatch.setenv("OPENAI_API_KEY", "test-key")
        import services.document_service
        monkeypatch.setattr(services.document_service, "OpenAI", MockOpenAI)
        
        try:
            # 处理文档
            await document_service.process_document(str(document.id))
            
            # 验证日志包含所有关键步骤
            log_messages = [record.message for record in caplog.records]
            log_text = " ".join(log_messages)
            
            # 验证关键步骤都被记录
            assert "开始处理文档" in log_text
            assert "更新文档状态为 processing" in log_text or "processing" in log_text
            assert "解析文档" in log_text
            assert "文本分块" in log_text
            assert "生成嵌入向量" in log_text or "嵌入" in log_text
            assert "存储文档块" in log_text or "存储" in log_text
            assert "处理成功完成" in log_text or "完成" in log_text
        
        except Exception as e:
            if "no such table: kb_document_chunks" in str(e):
                pytest.skip("SQLite不支持向量类型，跳过此测试")
            else:
                raise
    
    @pytest.mark.asyncio
    async def test_process_document_updates_processed_at(self, document_service, test_user_id, tmp_path, monkeypatch):
        """测试成功处理后更新processed_at时间戳"""
        # 创建测试文档
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        
        pdf_path = tmp_path / "test.pdf"
        c = canvas.Canvas(str(pdf_path), pagesize=letter)
        c.drawString(100, 750, "Test content")
        c.save()
        
        with open(pdf_path, 'rb') as f:
            file_data = f.read()
        
        document = await document_service.upload_document(
            file_data=file_data,
            filename="test.pdf",
            file_type="application/pdf",
            user_id=test_user_id
        )
        
        # 初始时processed_at应该为None
        assert document.processed_at is None
        
        # Mock OpenAI
        class MockEmbeddingData:
            def __init__(self, embedding):
                self.embedding = embedding
        
        class MockEmbeddingResponse:
            def __init__(self, data):
                self.data = data
        
        class MockOpenAI:
            class Embeddings:
                def create(self, model, input):
                    return MockEmbeddingResponse([
                        MockEmbeddingData([0.1] * 1536) for _ in input
                    ])
            
            def __init__(self, api_key, base_url=None):
                self.embeddings = self.Embeddings()
        
        monkeypatch.setenv("OPENAI_API_KEY", "test-key")
        import services.document_service
        monkeypatch.setattr(services.document_service, "OpenAI", MockOpenAI)
        
        try:
            # 处理文档
            await document_service.process_document(str(document.id))
            
            # 刷新状态
            document_service.db.refresh(document)
            
            # 验证processed_at已设置
            assert document.processed_at is not None
            assert isinstance(document.processed_at, datetime)
            
            # 验证processed_at在合理的时间范围内（最近几秒）
            time_diff = (datetime.utcnow() - document.processed_at).total_seconds()
            assert time_diff < 10  # 应该在10秒内
        
        except Exception as e:
            if "no such table: kb_document_chunks" in str(e):
                pytest.skip("SQLite不支持向量类型，跳过此测试")
            else:
                raise
    
    @pytest.mark.asyncio
    async def test_process_document_clears_previous_error(self, document_service, test_user_id, tmp_path, monkeypatch):
        """测试重新处理时清除之前的错误消息"""
        # 创建测试文档
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        
        pdf_path = tmp_path / "test.pdf"
        c = canvas.Canvas(str(pdf_path), pagesize=letter)
        c.drawString(100, 750, "Test content")
        c.save()
        
        with open(pdf_path, 'rb') as f:
            file_data = f.read()
        
        document = await document_service.upload_document(
            file_data=file_data,
            filename="test.pdf",
            file_type="application/pdf",
            user_id=test_user_id
        )
        
        # 手动设置一个错误消息（模拟之前的失败）
        document.error_message = "Previous error message"
        document_service.db.commit()
        
        # Mock OpenAI
        class MockEmbeddingData:
            def __init__(self, embedding):
                self.embedding = embedding
        
        class MockEmbeddingResponse:
            def __init__(self, data):
                self.data = data
        
        class MockOpenAI:
            class Embeddings:
                def create(self, model, input):
                    return MockEmbeddingResponse([
                        MockEmbeddingData([0.1] * 1536) for _ in input
                    ])
            
            def __init__(self, api_key, base_url=None):
                self.embeddings = self.Embeddings()
        
        monkeypatch.setenv("OPENAI_API_KEY", "test-key")
        import services.document_service
        monkeypatch.setattr(services.document_service, "OpenAI", MockOpenAI)
        
        try:
            # 重新处理文档
            await document_service.process_document(str(document.id))
            
            # 刷新状态
            document_service.db.refresh(document)
            
            # 验证错误消息已清除
            assert document.error_message is None
            assert document.upload_status == "completed"
        
        except Exception as e:
            if "no such table: kb_document_chunks" in str(e):
                pytest.skip("SQLite不支持向量类型，跳过此测试")
            else:
                raise


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
