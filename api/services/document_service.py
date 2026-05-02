"""
文档服务
负责知识库文档的上传、验证、解析、分块和向量。
"""
import os
import uuid
from datetime import datetime
from typing import List, Dict, Any, Optional
from sqlalchemy.orm import Session
from sqlalchemy import delete, select, func, text
import logging
import PyPDF2
from docx import Document as DocxDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
import tiktoken

from models.kb_document import KBDocument, KBDocumentChunk
from services.config_service import ConfigService
from services.kb_vector_schema_service import KBVectorSchemaService
from openai import OpenAI
import time

logger = logging.getLogger(__name__)


class DocumentService:
    """
    文档服务。
    处理知识库文档的上传、验证、存储和处理
    """
    
    def __init__(self, db: Session):
        """
        初始化文档服。
        
        Args:
            db: 数据库会。
        """
        self.db = db
        self.upload_dir = os.getenv("KB_UPLOAD_DIR", "uploads/kb_documents")
        self.max_file_size = int(os.getenv("KB_MAX_FILE_SIZE", str(10 * 1024 * 1024)))  # 10MB
        self.allowed_types = [
            "application/pdf",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ]
        
        # 文本分块配置
        self.chunk_size = int(os.getenv("KB_CHUNK_SIZE", "1000"))
        self.chunk_overlap = int(os.getenv("KB_CHUNK_OVERLAP", "200"))
        
        # 初始化文本分割器
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", "。", ".", "!", "?", " ", ""]
        )
        
        # 初始化tokenizer用于计算token数量
        try:
            self.tokenizer = tiktoken.encoding_for_model("gpt-3.5-turbo")
        except Exception as e:
            logger.warning(f"无法加载tokenizer，使用默认编。 {e}")
            self.tokenizer = tiktoken.get_encoding("cl100k_base")
        
        # 确保上传目录存在
        os.makedirs(self.upload_dir, exist_ok=True)

    def _is_postgresql(self) -> bool:
        bind = self.db.get_bind()
        return bind is not None and bind.dialect.name == "postgresql"

    @staticmethod
    def _parse_embedding_value(value: Any) -> Optional[List[float]]:
        if value is None:
            return None
        if isinstance(value, list):
            return [float(item) for item in value]

        text_value = str(value).strip()
        if not text_value:
            return []

        if text_value[0] in "[{" and text_value[-1] in "]}":
            text_value = text_value[1:-1].strip()

        if not text_value:
            return []

        return [float(item.strip()) for item in text_value.split(",") if item.strip()]

    def _build_chunk_record_from_row(self, row: Any) -> KBDocumentChunk:
        chunk = KBDocumentChunk(
            id=row.id,
            document_id=row.document_id,
            chunk_index=row.chunk_index,
            content=row.content,
            token_count=row.token_count,
            created_at=row.created_at,
        )
        embedding_value = getattr(row, "embedding_text", getattr(row, "embedding", None))
        chunk.embedding = self._parse_embedding_value(embedding_value)
        return chunk
    
    def validate_file_type(self, file_type: str) -> bool:
        """
        验证文件类型是否支持
        
        Args:
            file_type: MIME类型
            
        Returns:
            bool: 是否支持该文件类。
        """
        return file_type in self.allowed_types
    
    def validate_file_size(self, file_size: int) -> bool:
        """
        验证文件大小是否在限制内
        
        Args:
            file_size: 文件大小（字节）
            
        Returns:
            bool: 是否在大小限制内
        """
        return file_size <= self.max_file_size
    
    def generate_unique_filename(self, original_filename: str) -> str:
        """
        生成唯一的文件名以避免冲。
        
        Args:
            original_filename: 原始文件。
            
        Returns:
            str: 唯一的文件名
        """
        # 获取文件扩展。
        _, ext = os.path.splitext(original_filename)
        # 生成UUID作为文件。
        unique_name = f"{uuid.uuid4()}{ext}"
        return unique_name
    
    async def upload_document(
        self,
        file_data: bytes,
        filename: str,
        file_type: str,
        user_id: str
    ) -> KBDocument:
        """
        上传并验证文。
        
        Args:
            file_data: 文件二进制数。
            filename: 原始文件。
            file_type: MIME类型
            user_id: 上传用户ID
            
        Returns:
            KBDocument: 创建的文档记。
            
        Raises:
            ValueError: 文件类型或大小不符合要求
            IOError: 文件保存失败
        """
        # 验证文件类型
        if not self.validate_file_type(file_type):
            raise ValueError(
                f"不支持的文件类型，仅支持 PDF 。DOCX 格式。收到的类型: {file_type}"
            )
        
        # 验证文件大小
        file_size = len(file_data)
        if not self.validate_file_size(file_size):
            max_mb = self.max_file_size / (1024 * 1024)
            raise ValueError(
                f"文件大小超过限制（最。{max_mb:.0f}MB）。文件大。 {file_size / (1024 * 1024):.2f}MB"
            )
        
        # 生成唯一文件。
        unique_filename = self.generate_unique_filename(filename)
        file_path = os.path.join(self.upload_dir, unique_filename)
        
        try:
            # 保存文件到磁。
            with open(file_path, 'wb') as f:
                f.write(file_data)
            
            logger.info(f"文件已保。 {file_path}")
            
            # 创建数据库记。
            document = KBDocument(
                filename=filename,
                file_path=file_path,
                file_type=file_type,
                file_size=file_size,
                upload_status='pending',
                uploaded_by=uuid.UUID(user_id),
                uploaded_at=datetime.utcnow()
            )
            
            self.db.add(document)
            self.db.commit()
            self.db.refresh(document)
            
            logger.info(f"文档记录已创。 {document.id}, 文件。 {filename}")
            
            return document
            
        except IOError as e:
            logger.error(f"文件保存失败: {e}", exc_info=True)
            # 如果文件已创建，尝试删除
            if os.path.exists(file_path):
                try:
                    os.remove(file_path)
                except Exception:
                    pass
            raise IOError(f"文件保存失败，请稍后重试: {str(e)}")
        
        except Exception as e:
            logger.error(f"文档上传失败: {e}", exc_info=True)
            self.db.rollback()
            # 清理已保存的文件
            if os.path.exists(file_path):
                try:
                    os.remove(file_path)
                except Exception:
                    pass
            raise
    
    def list_documents(
        self,
        page: int = 1,
        page_size: int = 20
    ) -> Dict[str, Any]:
        """
        列出文档，支持分。
        
        Args:
            page: 页码（从1开始）
            page_size: 每页数量
            
        Returns:
            Dict: 包含文档列表、总数、页码等信息
        """
        try:
            # 计算偏移。
            offset = (page - 1) * page_size
            
            # 查询总数
            total = self.db.execute(
                select(func.count(KBDocument.id))
            ).scalar()
            
            # 查询文档列表
            documents = self.db.execute(
                select(KBDocument)
                .order_by(KBDocument.uploaded_at.desc())
                .offset(offset)
                .limit(page_size)
            ).scalars().all()
            
            return {
                "documents": documents,
                "total": total,
                "page": page,
                "page_size": page_size,
                "total_pages": (total + page_size - 1) // page_size if total > 0 else 0
            }
            
        except Exception as e:
            logger.error(f"列出文档失败: {e}", exc_info=True)
            raise
    
    async def delete_document(self, document_id: str) -> bool:
        """
        删除文档及其所有相关数。
        
        Args:
            document_id: 文档ID
            
        Returns:
            bool: 是否删除成功
            
        Raises:
            ValueError: 文档不存。
        """
        try:
            document_uuid = uuid.UUID(document_id)

            # Only load scalar metadata needed for file cleanup so delete
            # never triggers ORM loading of pgvector-backed chunk rows.
            file_path = self.db.execute(
                select(KBDocument.file_path).where(KBDocument.id == document_uuid)
            ).scalar_one_or_none()
            
            if not file_path:
                raise ValueError(f"文档不存。 {document_id}")

            delete_result = self.db.execute(
                delete(KBDocument).where(KBDocument.id == document_uuid)
            )
            if delete_result.rowcount == 0:
                self.db.rollback()
                raise ValueError(f"文档不存。 {document_id}")

            self.db.commit()
            
            # 删除文件
            if os.path.exists(file_path):
                try:
                    os.remove(file_path)
                    logger.info(f"文件已删。 {file_path}")
                except Exception as e:
                    logger.warning(f"删除文件失败: {file_path}, 错误: {e}")
            
            logger.info(f"文档已删。 {document_id}")
            return True
            
        except ValueError:
            raise
        except Exception as e:
            logger.error(f"删除文档失败: {e}", exc_info=True)
            self.db.rollback()
            raise
    
    async def parse_pdf(self, file_path: str) -> str:
        """
        解析PDF文档，提取文本内。
        
        Args:
            file_path: PDF文件路径
            
        Returns:
            str: 提取的文本内。
            
        Raises:
            ValueError: 文件损坏或无法解。
        """
        try:
            text_content = []
            
            with open(file_path, 'rb') as file:
                # 创建PDF读取。
                pdf_reader = PyPDF2.PdfReader(file)
                
                # 检查PDF是否加密
                if pdf_reader.is_encrypted:
                    logger.warning(f"PDF文件已加。 {file_path}")
                    raise ValueError("PDF文件已加密，无法解析")
                
                # 提取每一页的文本
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        # 提取文本内容
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text_content.append(page_text)
                    except Exception as e:
                        logger.warning(f"提取PDF第{page_num + 1}页失。 {e}")
                        # 继续处理其他页面
                        continue
            
            # 合并所有文。
            full_text = "\n\n".join(text_content)
            
            if not full_text.strip():
                logger.warning(f"PDF文件未提取到文本内容: {file_path}")
                raise ValueError("PDF文件未包含可提取的文本内容")
            
            logger.info(f"成功解析PDF文件: {file_path}, 提取文本长度: {len(full_text)}")
            return full_text
            
        except PyPDF2.errors.PdfReadError as e:
            logger.error(f"PDF文件损坏或格式错。 {file_path}, 错误: {e}")
            raise ValueError(f"PDF文件损坏或格式错。 {str(e)}")
        except FileNotFoundError:
            logger.error(f"文件不存。 {file_path}")
            raise ValueError(f"文件不存。 {file_path}")
        except Exception as e:
            logger.error(f"解析PDF文件失败: {file_path}, 错误: {e}", exc_info=True)
            raise ValueError(f"解析PDF文件失败: {str(e)}")
    
    async def parse_docx(self, file_path: str) -> str:
        """
        解析DOCX文档，提取文本内。
        
        Args:
            file_path: DOCX文件路径
            
        Returns:
            str: 提取的文本内。
            
        Raises:
            ValueError: 文件损坏或无法解。
        """
        # 首先检查文件是否存。
        if not os.path.exists(file_path):
            logger.error(f"文件不存。 {file_path}")
            raise ValueError(f"文件不存。 {file_path}")
        
        try:
            text_content = []
            
            # 打开DOCX文档
            doc = DocxDocument(file_path)
            
            # 提取段落文本
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    text_content.append(text)
            
            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            # 合并所有文。
            full_text = "\n\n".join(text_content)
            
            if not full_text.strip():
                logger.warning(f"DOCX文件未提取到文本内容: {file_path}")
                raise ValueError("DOCX文件未包含可提取的文本内容")
            
            logger.info(f"成功解析DOCX文件: {file_path}, 提取文本长度: {len(full_text)}")
            return full_text
            
        except Exception as e:
            logger.error(f"解析DOCX文件失败: {file_path}, 错误: {e}", exc_info=True)
            raise ValueError(f"解析DOCX文件失败: {str(e)}")
    
    async def parse_document(self, file_path: str, file_type: str) -> str:
        """
        根据文件类型解析文档，提取文本内。
        
        Args:
            file_path: 文件路径
            file_type: MIME类型
            
        Returns:
            str: 提取的文本内。
            
        Raises:
            ValueError: 不支持的文件类型或解析失。
        """
        try:
            if file_type == "application/pdf":
                return await self.parse_pdf(file_path)
            elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                return await self.parse_docx(file_path)
            else:
                raise ValueError(f"不支持的文件类型: {file_type}")
        except ValueError:
            # 重新抛出ValueError（包含用户友好的错误消息。
            raise
        except Exception as e:
            logger.error(f"解析文档失败: {file_path}, 类型: {file_type}, 错误: {e}", exc_info=True)
            raise ValueError(f"解析文档失败: {str(e)}")
    
    def count_tokens(self, text: str) -> int:
        """
        计算文本的token数量
        
        Args:
            text: 要计算的文本
            
        Returns:
            int: token数量
        """
        try:
            tokens = self.tokenizer.encode(text)
            return len(tokens)
        except Exception as e:
            logger.warning(f"计算token数量失败，使用字符数估算: {e}")
            # 如果tokenizer失败，使用粗略估算（平均每个token。个字符）
            return len(text) // 4
    
    async def chunk_text(self, text: str) -> List[Dict[str, Any]]:
        """
        将文本分割成语义。
        
        使用RecursiveCharacterTextSplitter在段落边界处分割文本。
        保持语义连贯性。每个块包含内容和token计数。
        
        Args:
            text: 要分割的文本内容
            
        Returns:
            List[Dict[str, Any]]: 分块列表，每个元素包。
                - content: 块的文本内容
                - token_count: 块的token数量
                - chunk_index: 块的索引（从0开始）
                
        Raises:
            ValueError: 文本为空或分块失。
        """
        try:
            # 验证输入
            if not text or not text.strip():
                raise ValueError("文本内容为空，无法分块")
            
            # 使用LangChain的RecursiveCharacterTextSplitter进行分块
            # 它会尝试在段落边界（\n\n）、句子边界（。！？）等处分割
            chunks = self.text_splitter.split_text(text)
            
            if not chunks:
                raise ValueError("文本分块失败，未生成任何块")
            
            # 为每个块添加元数。
            result = []
            for index, chunk_content in enumerate(chunks):
                # 计算token数量
                token_count = self.count_tokens(chunk_content)
                
                result.append({
                    "content": chunk_content,
                    "token_count": token_count,
                    "chunk_index": index
                })
            
            logger.info(
                f"文本分块完成: 原始长度={len(text)}, "
                f"块数。{len(result)}, "
                f"平均token。{sum(c['token_count'] for c in result) / len(result):.1f}"
            )
            
            return result
            
        except ValueError:
            # 重新抛出ValueError
            raise
        except Exception as e:
            logger.error(f"文本分块失败: {e}", exc_info=True)
            raise ValueError(f"文本分块失败: {str(e)}")
    
    async def generate_embeddings(self, texts: List[str]) -> List[List[float]]:
        """
        使用OpenAI生成文本嵌入向量

        支持批处理（每批100个文本）和重试逻辑。次重试，指数退避）
        优雅处理速率限制和API故障

        Args:
            texts: 要生成嵌入的文本列表

        Returns:
            List[List[float]]: 嵌入向量列表

        Raises:
            ValueError: 文本列表为空或API配置无效
            RuntimeError: API调用失败（重。次后。
        """
        try:
            # 验证输入
            if not texts:
                raise ValueError("文本列表为空，无法生成嵌入")

            # 获取OpenAI API配置
            # 注意：这里需要从环境变量或配置中获取API密钥
            # 获取向量配置
            config_service = ConfigService(self.db)
            vector_config = await config_service.get_vector_config()
            
            if not vector_config.api_key:
                raise ValueError("向量模型API密钥未配置")

            # 初始化OpenAI客户端（使用向量配置）
            base_url = vector_config.api_endpoint
            if base_url.endswith('/embeddings'):
                base_url = base_url[:-len('/embeddings')]
            
            client = OpenAI(api_key=vector_config.api_key, base_url=base_url)

            # 使用配置的嵌入模型
            embedding_model = vector_config.model_name
            batch_size = 100  # 每批处理100个文。
            max_retries = 3  # 最大重试次。

            all_embeddings = []

            # 分批处理文本
            for i in range(0, len(texts), batch_size):
                batch = texts[i:i + batch_size]
                batch_num = i // batch_size + 1
                total_batches = (len(texts) + batch_size - 1) // batch_size

                logger.info(
                    f"处理嵌入批次 {batch_num}/{total_batches}, "
                    f"文本数量: {len(batch)}"
                )

                # 重试逻辑：指数退。
                retry_count = 0
                last_error = None

                while retry_count < max_retries:
                    try:
                        # 调用OpenAI API生成嵌入
                        response = client.embeddings.create(
                            model=embedding_model,
                            input=batch
                        )

                        # 提取嵌入向量
                        batch_embeddings = [item.embedding for item in response.data]
                        all_embeddings.extend(batch_embeddings)

                        logger.info(
                            f"批次 {batch_num}/{total_batches} 嵌入生成成功, "
                            f"向量维度: {len(batch_embeddings[0]) if batch_embeddings else 0}"
                        )

                        # 成功，跳出重试循。
                        break

                    except Exception as e:
                        retry_count += 1
                        last_error = e

                        # 检查是否是速率限制错误
                        error_str = str(e).lower()
                        is_rate_limit = (
                            "rate limit" in error_str or
                            "rate_limit" in error_str or
                            "429" in error_str
                        )

                        if retry_count < max_retries:
                            # 计算退避时间：指数退。
                            # 速率限制：等待更长时。
                            if is_rate_limit:
                                wait_time = min(2 ** retry_count * 5, 60)  # 5s, 10s, 20s (最多60s)
                                logger.warning(
                                    f"遇到速率限制，批次 {batch_num}/{total_batches}, "
                                    f"等待 {wait_time}秒后重试 (尝试 {retry_count}/{max_retries})"
                                )
                            else:
                                wait_time = 2 ** retry_count  # 1s, 2s, 4s
                                logger.warning(
                                    f"API调用失败，批次 {batch_num}/{total_batches}, "
                                    f"错误: {e}, 等待 {wait_time}秒后重试 (尝试 {retry_count}/{max_retries})"
                                )

                            time.sleep(wait_time)
                        else:
                            # 达到最大重试次数
                            logger.error(
                                f"批次 {batch_num}/{total_batches} 嵌入生成失败，"
                                f"已重试 {max_retries} 次，错误: {last_error}",
                                exc_info=True
                            )
                            raise RuntimeError(
                                f"嵌入生成失败（批次 {batch_num}/{total_batches}），"
                                f"已重试 {max_retries} 次，错误: {str(last_error)}"
                            )

            logger.info(
                f"所有嵌入生成完成，总文本数={len(texts)}, "
                f"总向量数={len(all_embeddings)}, "
                f"批次数={total_batches}"
            )

            return all_embeddings

        except ValueError:
            # 重新抛出ValueError（输入验证错误）
            raise
        except RuntimeError:
            # 重新抛出RuntimeError（API调用失败。
            raise
        except Exception as e:
            logger.error(f"生成嵌入失败: {e}", exc_info=True)
            raise RuntimeError(f"生成嵌入失败: {str(e)}")
    

    
    async def store_chunks_with_embeddings(
        self,
        document_id: str,
        chunks: List[Dict[str, Any]],
        embeddings: List[List[float]]
    ) -> List[KBDocumentChunk]:
        """
        将文档块及其嵌入向量存储到数据库
        
        Args:
            document_id: 文档ID
            chunks: 文档块列表，每个包含 content, token_count, chunk_index
            embeddings: 嵌入向量列表，与chunks一一对应
            
        Returns:
            List[KBDocumentChunk]: 创建的文档块记录列表
            
        Raises:
            ValueError: 输入验证失败
            RuntimeError: 数据库操作失败
        """
        try:
            # 验证输入
            if not chunks:
                raise ValueError("文档块列表为空")
            
            if len(chunks) != len(embeddings):
                raise ValueError(
                    f"文档块数量({len(chunks)})与嵌入向量数量({len(embeddings)})不匹配"
                )
            
            # 获取向量配置以验证维度
            config_service = ConfigService(self.db)
            vector_config = await config_service.get_vector_config()
            expected_dim = vector_config.embedding_dimension or 1536
            KBVectorSchemaService.ensure_schema_matches_dimension(
                db=self.db,
                target_dimension=int(expected_dim),
            )
            
            # 验证文档存在
            document = self.db.execute(
                select(KBDocument).where(KBDocument.id == uuid.UUID(document_id))
            ).scalar_one_or_none()
            
            if not document:
                raise ValueError(f"文档不存在: {document_id}")
            
            # 创建文档块记录
            chunk_records = []
            for chunk_data, embedding in zip(chunks, embeddings):
                # 验证嵌入向量维度
                if len(embedding) != expected_dim:
                    raise ValueError(
                        f"嵌入向量维度错误: 期望{expected_dim}，实际{len(embedding)}"
                    )
                
                chunk_record = KBDocumentChunk(
                    id=uuid.uuid4(),
                    document_id=uuid.UUID(document_id),
                    chunk_index=chunk_data["chunk_index"],
                    content=chunk_data["content"],
                    token_count=chunk_data["token_count"],
                    embedding=embedding,
                    created_at=datetime.utcnow()
                )
                chunk_records.append(chunk_record)
            
            # 批量插入数据库
            self.db.add_all(chunk_records)
            self.db.commit()
            
            # 刷新记录以获取生成的ID
            logger.info(
                f"成功存储 {len(chunk_records)} 个文档块到数据库, "
                f"文档ID: {document_id}"
            )
            
            return chunk_records
            
        except ValueError:
            # 重新抛出验证错误
            raise
        except Exception as e:
            logger.error(
                f"存储文档块失败: 文档ID={document_id}, 错误={e}",
                exc_info=True
            )
            self.db.rollback()
            if "expected" in str(e) and "dimensions" in str(e):
                raise RuntimeError(
                    "知识库向量维度与当前配置不一致，请同步向量配置后重试"
                )
            raise RuntimeError(f"存储文档块失败: {str(e)}")
    
    async def retrieve_chunks_by_document(
        self,
        document_id: str
    ) -> List[KBDocumentChunk]:
        """
        检索指定文档的所有块（包含嵌入向量）
        
        Args:
            document_id: 文档ID
            
        Returns:
            List[KBDocumentChunk]: 文档块列表，按chunk_index排序
            
        Raises:
            ValueError: 文档不存在
        """
        try:
            # 验证文档存在
            document = self.db.execute(
                select(KBDocument).where(KBDocument.id == uuid.UUID(document_id))
            ).scalar_one_or_none()
            
            if not document:
                raise ValueError(f"文档不存在: {document_id}")
            
            # 查询文档块，按chunk_index排序
            if self._is_postgresql():
                rows = self.db.execute(
                    text(
                        """
                        SELECT
                            id,
                            document_id,
                            chunk_index,
                            content,
                            token_count,
                            embedding::text AS embedding_text,
                            created_at
                        FROM kb_document_chunks
                        WHERE document_id = :document_id
                        ORDER BY chunk_index
                        """
                    ),
                    {"document_id": uuid.UUID(document_id)},
                ).fetchall()
                chunks = [self._build_chunk_record_from_row(row) for row in rows]
            else:
                chunks = self.db.execute(
                    select(KBDocumentChunk)
                    .where(KBDocumentChunk.document_id == uuid.UUID(document_id))
                    .order_by(KBDocumentChunk.chunk_index)
                ).scalars().all()
            
            logger.info(
                f"检索到 {len(chunks)} 个文档块, 文档ID: {document_id}"
            )
            
            return chunks
            
        except ValueError:
            raise
        except Exception as e:
            logger.error(
                f"检索文档块失败: 文档ID={document_id}, 错误={e}",
                exc_info=True
            )
            raise RuntimeError(f"检索文档块失败: {str(e)}")
    
    async def retrieve_chunk_by_id(
        self,
        chunk_id: str
    ) -> Optional[KBDocumentChunk]:
        """
        根据ID检索单个文档块（包含嵌入向量）
        
        Args:
            chunk_id: 文档块ID
            
        Returns:
            Optional[KBDocumentChunk]: 文档块记录，如果不存在则返回None
        """
        try:
            if self._is_postgresql():
                row = self.db.execute(
                    text(
                        """
                        SELECT
                            id,
                            document_id,
                            chunk_index,
                            content,
                            token_count,
                            embedding::text AS embedding_text,
                            created_at
                        FROM kb_document_chunks
                        WHERE id = :chunk_id
                        """
                    ),
                    {"chunk_id": uuid.UUID(chunk_id)},
                ).fetchone()
                chunk = self._build_chunk_record_from_row(row) if row else None
            else:
                chunk = self.db.execute(
                    select(KBDocumentChunk)
                    .where(KBDocumentChunk.id == uuid.UUID(chunk_id))
                ).scalar_one_or_none()
            
            if chunk:
                logger.info(f"检索到文档块: {chunk_id}")
            else:
                logger.warning(f"文档块不存在: {chunk_id}")
            
            return chunk
            
        except Exception as e:
            logger.error(
                f"检索文档块失败: 块ID={chunk_id}, 错误={e}",
                exc_info=True
            )
            raise RuntimeError(f"检索文档块失败: {str(e)}")
    
    async def process_document(self, document_id: str) -> None:
        """
        异步处理文档：解析、分块、生成嵌入、存储向量

        这是文档处理的主要编排方法，协调以下步骤：
        1. 更新状态为 "processing"
        2. 解析文档提取文本
        3. 将文本分块
        4. 生成嵌入向量
        5. 存储块和向量到数据库
        6. 更新状态为 "completed"
        
        如果任何步骤失败，将回滚部分数据并更新状态为 "failed"
        
        Args:
            document_id: 要处理的文档ID
            
        Raises:
            ValueError: 文档不存在
            RuntimeError: 处理失败
        """
        document = None
        chunks_created = False
        document_uuid = uuid.UUID(document_id)

        def delete_existing_chunks() -> int:
            deleted_count = (
                self.db.query(KBDocumentChunk)
                .filter(KBDocumentChunk.document_id == document_uuid)
                .delete(synchronize_session=False)
            )
            if deleted_count:
                self.db.commit()
                logger.info(
                    "重处理前已清理 %s 个旧文档块: %s",
                    deleted_count,
                    document_id,
                )
            return deleted_count

        def mark_failed(error_message: str) -> None:
            if document is None:
                return

            try:
                if chunks_created:
                    delete_existing_chunks()

                document.upload_status = "failed"
                document.error_message = error_message
                self.db.commit()
                logger.info(f"文档状态已更新为 failed: {document_id}")
            except Exception as rollback_error:
                logger.error(
                    f"回滚失败: {document_id}, 错误: {rollback_error}",
                    exc_info=True
                )
                self.db.rollback()

        try:
            # 1. 获取文档记录
            logger.info(f"开始处理文档: {document_id}")

            document = self.db.execute(
                select(KBDocument).where(KBDocument.id == document_uuid)
            ).scalar_one_or_none()

            if not document:
                raise ValueError(f"文档不存在: {document_id}")

            # 2. 更新状态为 "processing"
            logger.info(f"更新文档状态为 processing: {document_id}")
            document.upload_status = "processing"
            document.processed_at = None
            document.error_message = None
            self.db.commit()
            self.db.refresh(document)

            # 3. 解析文档提取文本
            logger.info(f"解析文档: {document.filename} ({document.file_type})")
            text_content = await self.parse_document(
                document.file_path,
                document.file_type
            )
            logger.info(f"文档解析完成，提取文本长度: {len(text_content)}")

            # 4. 将文本分块
            logger.info(f"开始文本分块: {document_id}")
            chunks = await self.chunk_text(text_content)
            logger.info(f"文本分块完成，生成 {len(chunks)} 个块")

            # 5. 生成嵌入向量
            logger.info(f"开始生成嵌入向量: {document_id}")
            chunk_texts = [chunk["content"] for chunk in chunks]
            embeddings = await self.generate_embeddings(chunk_texts)
            logger.info(f"嵌入向量生成完成，共 {len(embeddings)} 个向量")

            # 6. 重处理前先清理旧块，避免唯一键冲突
            delete_existing_chunks()

            # 7. 存储块和向量到数据库
            logger.info(f"开始存储文档块和向量: {document_id}")
            chunk_records = await self.store_chunks_with_embeddings(
                str(document.id),
                chunks,
                embeddings
            )
            chunks_created = True
            logger.info(f"文档块存储完成，共 {len(chunk_records)} 个块")

            # 8. 更新状态为 "completed"
            document.upload_status = "completed"
            document.processed_at = datetime.utcnow()
            document.error_message = None
            self.db.commit()
            self.db.refresh(document)

            logger.info(
                f"文档处理成功完成: {document_id}, "
                f"文件名: {document.filename}, "
                f"块数: {len(chunk_records)}"
            )

        except ValueError as e:
            error_msg = str(e)
            logger.error(f"文档处理失败（验证错误）: {document_id}, 错误: {error_msg}")
            mark_failed(error_msg)
            raise

        except RuntimeError as e:
            error_msg = str(e)
            logger.error(
                f"文档处理失败（运行时错误）: {document_id}, 错误: {error_msg}",
                exc_info=True
            )
            mark_failed(error_msg)
            raise

        except Exception as e:
            error_msg = f"未预期的错误: {str(e)}"
            logger.error(
                f"文档处理失败（未预期错误）: {document_id}, 错误: {e}",
                exc_info=True
            )
            mark_failed(error_msg)
            raise RuntimeError(error_msg)

    async def process_document(self, document_id: str) -> None:
        """Final document-processing implementation with safe reprocessing support."""
        document = None
        chunks_created = False

        def mark_failed(error_message: str) -> None:
            if document is None:
                return

            try:
                if chunks_created:
                    self.db.query(KBDocumentChunk).filter(
                        KBDocumentChunk.document_id == uuid.UUID(document_id)
                    ).delete(synchronize_session=False)
                    self.db.commit()

                document.upload_status = "failed"
                document.error_message = error_message
                self.db.commit()
            except Exception as rollback_error:
                logger.error(
                    f"回滚失败: {document_id}, 错误: {rollback_error}",
                    exc_info=True,
                )
                self.db.rollback()

        try:
            logger.info(f"开始处理文档: {document_id}")

            document = self.db.execute(
                select(KBDocument).where(KBDocument.id == uuid.UUID(document_id))
            ).scalar_one_or_none()

            if not document:
                raise ValueError(f"文档不存在: {document_id}")

            document.upload_status = "processing"
            document.error_message = None
            self.db.commit()
            self.db.refresh(document)

            logger.info(f"解析文档: {document.filename} ({document.file_type})")
            text_content = await self.parse_document(
                document.file_path,
                document.file_type,
            )

            logger.info(f"开始文本分块: {document_id}")
            chunks = await self.chunk_text(text_content)

            logger.info(f"开始生成嵌入向量: {document_id}")
            chunk_texts = [chunk["content"] for chunk in chunks]
            embeddings = await self.generate_embeddings(chunk_texts)

            existing_chunk_count = (
                self.db.query(KBDocumentChunk)
                .filter(KBDocumentChunk.document_id == uuid.UUID(document_id))
                .delete(synchronize_session=False)
            )
            if existing_chunk_count:
                self.db.commit()
                logger.info(
                    "Deleted %s existing chunks before reprocessing document %s",
                    existing_chunk_count,
                    document_id,
                )

            logger.info(f"开始存储文档块和向量: {document_id}")
            chunk_records = await self.store_chunks_with_embeddings(
                str(document.id),
                chunks,
                embeddings,
            )
            chunks_created = True

            document.upload_status = "completed"
            document.processed_at = datetime.utcnow()
            document.error_message = None
            self.db.commit()
            self.db.refresh(document)

            logger.info(
                f"文档处理成功完成: {document_id}, "
                f"文件名: {document.filename}, "
                f"块数: {len(chunk_records)}"
            )

        except ValueError as e:
            error_msg = str(e)
            logger.error(
                f"文档处理失败（验证错误）: {document_id}, 错误: {error_msg}"
            )
            mark_failed(error_msg)
            raise

        except RuntimeError as e:
            error_msg = str(e)
            logger.error(
                f"文档处理失败（运行时错误）: {document_id}, 错误: {error_msg}",
                exc_info=True,
            )
            mark_failed(error_msg)
            raise

        except Exception as e:
            error_msg = f"未预期的错误: {str(e)}"
            logger.error(
                f"文档处理失败（未预期错误）: {document_id}, 错误: {e}",
                exc_info=True,
            )
            mark_failed(error_msg)
            raise RuntimeError(error_msg)
    
    async def search_similar_chunks(
        self,
        query_embedding: List[float],
        top_k: int = 5,
        similarity_threshold: float = 0.0
    ) -> List[Dict[str, Any]]:
        """
        使用向量相似度搜索最相关的文档块
        
        使用pgvector的余弦相似度（cosine similarity）进行向量搜索。
        返回与查询向量最相似的top_k个文档块。
        
        Args:
            query_embedding: 查询的嵌入向量（1536维）
            top_k: 返回的最相似块数量（默认5）
            similarity_threshold: 最小相似度阈值（默认0.0，范围-1到1）
            
        Returns:
            List[Dict[str, Any]]: 相似文档块列表，每个包含：
                - chunk_id: 块ID
                - document_id: 文档ID
                - document_name: 文档名称
                - content: 块内容
                - chunk_index: 块索引
                - similarity_score: 相似度分数（1 - cosine_distance）
                
        Raises:
            ValueError: 输入验证失败
            RuntimeError: 数据库查询失败
        """
        try:
            # 验证输入
            if not query_embedding:
                raise ValueError("查询嵌入向量为空")
            
            # 获取向量配置以验证维度
            config_service = ConfigService(self.db)
            vector_config = await config_service.get_vector_config()
            expected_dim = vector_config.embedding_dimension or 1536
            KBVectorSchemaService.ensure_schema_matches_dimension(
                db=self.db,
                target_dimension=int(expected_dim),
            )
            
            if len(query_embedding) != expected_dim:
                raise ValueError(
                    f"查询嵌入向量维度错误: 期望{expected_dim}，实际{len(query_embedding)}"
                )
            
            if top_k <= 0:
                raise ValueError(f"top_k必须大于0，实际值: {top_k}")
            
            # 将Python列表转换为PostgreSQL数组格式
            # pgvector使用 <=> 操作符计算余弦距离（cosine distance）
            # 余弦距离 = 1 - 余弦相似度
            # 因此相似度 = 1 - 距离
            
            # 构建原始SQL查询以使用pgvector的向量操作符
            # 注意：SQLAlchemy不直接支持pgvector操作符，需要使用text()
            from sqlalchemy import text
            
            # 格式化向量为PostgreSQL数组字符串
            vector_str = "[" + ",".join(str(x) for x in query_embedding) + "]"
            
            # 使用 <=> 操作符计算余弦距离
            # ORDER BY embedding <=> query_vector 会按距离升序排列（最相似的在前）
            query = text("""
                SELECT 
                    c.id as chunk_id,
                    c.document_id,
                    c.content,
                    c.chunk_index,
                    c.token_count,
                    d.filename as document_name,
                    1 - (c.embedding <=> :query_vector::vector) as similarity_score
                FROM kb_document_chunks c
                JOIN kb_documents d ON c.document_id = d.id
                WHERE d.upload_status = 'completed'
                    AND 1 - (c.embedding <=> :query_vector::vector) >= :threshold
                ORDER BY c.embedding <=> :query_vector::vector
                LIMIT :limit
            """)
            
            # 执行查询
            result = self.db.execute(
                query,
                {
                    "query_vector": vector_str,
                    "threshold": similarity_threshold,
                    "limit": top_k
                }
            )
            
            # 处理结果
            similar_chunks = []
            for row in result:
                similar_chunks.append({
                    "chunk_id": str(row.chunk_id),
                    "document_id": str(row.document_id),
                    "document_name": row.document_name,
                    "content": row.content,
                    "chunk_index": row.chunk_index,
                    "token_count": row.token_count,
                    "similarity_score": float(row.similarity_score)
                })
            
            logger.info(
                f"向量相似度搜索完成: 找到 {len(similar_chunks)} 个相似块, "
                f"top_k={top_k}, threshold={similarity_threshold}"
            )
            
            # 如果找到结果，记录最高和最低相似度
            if similar_chunks:
                max_score = similar_chunks[0]["similarity_score"]
                min_score = similar_chunks[-1]["similarity_score"]
                logger.info(
                    f"相似度范围: {min_score:.4f} - {max_score:.4f}"
                )
            
            return similar_chunks
            
        except ValueError:
            # 重新抛出验证错误
            raise
        except Exception as e:
            logger.error(
                f"向量相似度搜索失败: {e}",
                exc_info=True
            )
            raise RuntimeError(f"向量相似度搜索失败: {str(e)}")

    async def process_document(self, document_id: str) -> None:
        """
        异步处理文档：解析、分块、生成嵌入、存储向量

        这是文档处理的主要编排方法，协调以下步骤：
        1. 更新状态为 "processing"
        2. 解析文档提取文本
        3. 将文本分块
        4. 生成嵌入向量
        5. 清理旧块并写入新块
        6. 更新状态为 "completed"

        如果任何步骤失败，将回滚部分数据并更新状态为 "failed"

        Args:
            document_id: 要处理的文档ID

        Raises:
            ValueError: 文档不存在
            RuntimeError: 处理失败
        """
        document = None
        chunks_created = False
        document_uuid = uuid.UUID(document_id)

        def delete_existing_chunks() -> int:
            deleted_count = (
                self.db.query(KBDocumentChunk)
                .filter(KBDocumentChunk.document_id == document_uuid)
                .delete(synchronize_session=False)
            )
            if deleted_count:
                self.db.commit()
                logger.info(
                    "重处理前已清理 %s 个旧文档块: %s",
                    deleted_count,
                    document_id,
                )
            return deleted_count

        def mark_failed(error_message: str) -> None:
            if document is None:
                return

            try:
                if chunks_created:
                    delete_existing_chunks()

                document.upload_status = "failed"
                document.error_message = error_message
                self.db.commit()
                logger.info(f"文档状态已更新为 failed: {document_id}")
            except Exception as rollback_error:
                logger.error(
                    f"回滚失败: {document_id}, 错误: {rollback_error}",
                    exc_info=True
                )
                self.db.rollback()

        try:
            # 1. 获取文档记录
            logger.info(f"开始处理文档: {document_id}")

            document = self.db.execute(
                select(KBDocument).where(KBDocument.id == document_uuid)
            ).scalar_one_or_none()

            if not document:
                raise ValueError(f"文档不存在: {document_id}")

            # 2. 更新状态为 "processing"
            logger.info(f"更新文档状态为 processing: {document_id}")
            document.upload_status = "processing"
            document.processed_at = None
            document.error_message = None
            self.db.commit()
            self.db.refresh(document)

            # 3. 解析文档提取文本
            logger.info(f"解析文档: {document.filename} ({document.file_type})")
            text_content = await self.parse_document(
                document.file_path,
                document.file_type
            )
            logger.info(f"文档解析完成，提取文本长度: {len(text_content)}")

            # 4. 将文本分块
            logger.info(f"开始文本分块: {document_id}")
            chunks = await self.chunk_text(text_content)
            logger.info(f"文本分块完成，生成 {len(chunks)} 个块")

            # 5. 生成嵌入向量
            logger.info(f"开始生成嵌入向量: {document_id}")
            chunk_texts = [chunk["content"] for chunk in chunks]
            embeddings = await self.generate_embeddings(chunk_texts)
            logger.info(f"嵌入向量生成完成，共 {len(embeddings)} 个向量")

            # 6. 重处理前清理旧块，避免 document_id + chunk_index 唯一键冲突
            delete_existing_chunks()

            # 7. 存储块和向量到数据库
            logger.info(f"开始存储文档块和向量: {document_id}")
            chunk_records = await self.store_chunks_with_embeddings(
                str(document.id),
                chunks,
                embeddings
            )
            chunks_created = True
            logger.info(f"文档块存储完成，共 {len(chunk_records)} 个块")

            # 8. 更新状态为 "completed"
            document.upload_status = "completed"
            document.processed_at = datetime.utcnow()
            document.error_message = None
            self.db.commit()
            self.db.refresh(document)

            logger.info(
                f"文档处理成功完成: {document_id}, "
                f"文件名: {document.filename}, "
                f"块数: {len(chunk_records)}"
            )

        except ValueError as e:
            error_msg = str(e)
            logger.error(f"文档处理失败（验证错误）: {document_id}, 错误: {error_msg}")
            mark_failed(error_msg)
            raise

        except RuntimeError as e:
            error_msg = str(e)
            logger.error(
                f"文档处理失败（运行时错误）: {document_id}, 错误: {error_msg}",
                exc_info=True
            )
            mark_failed(error_msg)
            raise

        except Exception as e:
            error_msg = f"未预期的错误: {str(e)}"
            logger.error(
                f"文档处理失败（未预期错误）: {document_id}, 错误: {e}",
                exc_info=True
            )
            mark_failed(error_msg)
            raise RuntimeError(error_msg)

